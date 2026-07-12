# -*- coding: utf-8 -*-
"""
身份证识别工具 v2.5 - PaddleOCR 2.x 稳定版本
- 支持身份证正面检测
- 优先提取"公民身份号码"标签后的号码
- 完整的OCR识别结果日志
- 文件拖拽功能
"""

import sys
import os
import re
import zipfile
import rarfile
import py7zr
import traceback
import tempfile
import shutil
from datetime import datetime
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                             QLabel, QPushButton, QTextEdit, QProgressBar, QFileDialog,
                             QMessageBox, QCheckBox, QListWidget, QAbstractItemView)
from PyQt5.QtCore import QThread, pyqtSignal, Qt
from pdfplumber import open as pdf_open
from PIL import Image
import docx

# PyMuPDF (fitz) 用于从PDF提取图片，可选依赖
try:
    import fitz
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False
    fitz = None

log_content = ""
error_log_path = os.path.join(os.path.expanduser("~"), "id_verifier_error.log")

def log_error(msg, exc_info=None):
    global log_content
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_content += f"[{timestamp}] ERROR: {msg}\n"
    if exc_info:
        log_content += "".join(traceback.format_exception(*exc_info))
    log_content += "--------------------------------------------------\n\n"
    try:
        with open(error_log_path, "w", encoding="utf-8") as f:
            f.write(log_content)
    except:
        pass

def find_id_numbers(text, is_id_card_front=False):
    if is_id_card_front:
        pattern2 = r'公民身份号码[：:]\s*([1-9]\d{5}(?:18|19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx])'
        matches = re.findall(pattern2, text, re.IGNORECASE)
        if matches:
            return list(set(matches))
    pattern = r'[1-9]\d{5}(?:18|19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx]'
    return list(set(re.findall(pattern, text, re.IGNORECASE)))

def is_id_card_front(text):
    text_lower = text.lower()
    front_keywords = ['公民', '身份', '号码', '姓名', '性别', '民族', '出生', '住址']
    back_keywords = ['签发', '机关', '有效期', '有效期限']
    front_score = sum(1 for kw in front_keywords if kw in text_lower)
    back_score = sum(1 for kw in back_keywords if kw in text_lower)
    return front_score >= 3 and back_score < 2

def extract_id_and_text_from_image(image_path):
    text = ""
    try:
        # 检查并转换GIF等特殊格式
        ext = os.path.splitext(image_path)[1].lower()
        if ext == '.gif':
            # 使用PIL打开并转换为PNG进行处理
            with Image.open(image_path) as img:
                # 使用tempfile创建安全的临时文件
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    temp_png = tmp.name
                try:
                    img.save(temp_png, 'PNG')
                    text = extract_text_from_image(temp_png, save_full_result=True)
                finally:
                    if os.path.exists(temp_png):
                        os.remove(temp_png)
        else:
            text = extract_text_from_image(image_path, save_full_result=True)
    except Exception as e:
        log_error(f"图片处理失败: {image_path}, 错误: {e}")
        text = ""

    log_error(f"\n处理文件: {os.path.basename(image_path)}")

    is_front = is_id_card_front(text)
    log_error(f"是否为身份证正面: {'是' if is_front else '否'}")

    if is_front:
        ids = find_id_numbers(text, is_id_card_front=True)
        if ids:
            log_error(f"从'公民身份号码'标签提取到: {ids}")
            return ids, text, True

    ids = find_id_numbers(text, is_id_card_front=False)
    log_error(f"提取到的身份证号码: {ids if ids else '无'}")

    return ids, text, is_front

def find_names(text):
    patterns = [
        r'姓名[：:\s]*([\u4e00-\u9fa5]{2,4})',
        r'名字[：:\s]*([\u4e00-\u9fa5]{2,4})',
    ]
    names = []
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        names.extend(matches)
    return list(set(names))

def is_hidden_name(name):
    return '*' in name or '×' in name or '不公布' in name

def filter_hidden_names(names):
    return [n for n in names if not is_hidden_name(n)]

def extract_text_from_pdf(pdf_path):
    try:
        with pdf_open(pdf_path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text
    except Exception as e:
        log_error(f"提取PDF文本失败: {pdf_path}, 错误: {e}")
        return ""

def extract_images_from_pdf(pdf_path, output_dir):
    images = []
    if not HAS_FITZ:
        log_error(f"未安装PyMuPDF，跳过PDF图片提取: {pdf_path}")
        return images
    try:
        doc = fitz.open(pdf_path)
        for page_num, page in enumerate(doc):
            for img_idx, img in enumerate(page.get_images()):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                if pix.n >= 5:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img_path = os.path.join(output_dir, f"pdf_img_{page_num}_{img_idx}.png")
                pix.save(img_path)
                images.append(img_path)
        doc.close()
    except Exception as e:
        log_error(f"提取PDF图片失败: {pdf_path}, 错误: {e}")
    return images

def extract_text_from_docx(docx_path):
    try:
        doc = docx.Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        return text
    except Exception as e:
        log_error(f"提取Word文本失败: {docx_path}, 错误: {e}")
        return ""

def extract_images_from_docx(docx_path, output_dir):
    images = []
    try:
        doc = docx.Document(docx_path)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_part = doc.part.related_parts[rel.target_ref]
                img_bytes = img_part.blob
                img_ext = os.path.splitext(rel.target_ref)[1] or '.png'
                img_path = os.path.join(output_dir, f"docx_img_{len(images)}{img_ext}")
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)
                images.append(img_path)
    except Exception as e:
        log_error(f"提取Word图片失败: {docx_path}, 错误: {e}")
    return images

def is_safe_path(base_dir, target_path):
    """检查路径是否安全，防止Zip Slip漏洞"""
    base_dir = os.path.abspath(base_dir)
    target_path = os.path.abspath(target_path)
    return os.path.commonpath([base_dir, target_path]) == base_dir

def process_archive(archive_path, temp_dir):
    extracted_dir = os.path.join(temp_dir, "extracted")
    os.makedirs(extracted_dir, exist_ok=True)
    try:
        if archive_path.lower().endswith('.zip'):
            with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                # 安全检查每个文件
                for member in zip_ref.infolist():
                    target_path = os.path.join(extracted_dir, member.filename)
                    if not is_safe_path(extracted_dir, target_path):
                        log_error(f"检测到不安全的文件路径: {member.filename}")
                        return None
                zip_ref.extractall(extracted_dir)
        elif archive_path.lower().endswith('.rar'):
            try:
                with rarfile.RarFile(archive_path, 'r') as rar_ref:
                    # 安全检查每个文件
                    for member in rar_ref.infolist():
                        target_path = os.path.join(extracted_dir, member.filename)
                        if not is_safe_path(extracted_dir, target_path):
                            log_error(f"检测到不安全的文件路径: {member.filename}")
                            return None
                    rar_ref.extractall(extracted_dir)
            except rarfile.RarCannotExec:
                log_error(f"RAR解压需要安装WinRAR或UnRAR工具!")
                log_error("请下载安装WinRAR: https://www.win-rar.com/download.html")
                log_error("或建议将RAR文件转换为ZIP格式")
                return None
        elif archive_path.lower().endswith('.7z'):
            with py7zr.SevenZipFile(archive_path, 'r') as sz:
                # 安全检查每个文件
                archive_contents = sz.getnames()
                for filename in archive_contents:
                    target_path = os.path.join(extracted_dir, filename)
                    if not is_safe_path(extracted_dir, target_path):
                        log_error(f"检测到不安全的文件路径: {filename}")
                        return None
                sz.extractall(extracted_dir)
    except Exception as e:
        log_error(f"解压文件失败: {archive_path}, 错误: {e}")
        return None
    log_error(f"解压成功: {archive_path} -> {extracted_dir}")
    return extracted_dir

def get_all_files(directory):
    files = []
    if not directory or not os.path.exists(directory):
        log_error(f"目录不存在: {directory}")
        return []
    for root, _, filenames in os.walk(directory):
        for filename in filenames:
            file_path = os.path.join(root, filename)
            files.append(file_path)
            log_error(f"发现文件: {file_path}")
    log_error(f"共找到 {len(files)} 个文件")
    return files

_paddle_ocr = None

def get_paddle_ocr():
    global _paddle_ocr
    if _paddle_ocr is None:
        from paddleocr import PaddleOCR
        script_dir = os.path.dirname(os.path.abspath(__file__))
        det_model_dir = os.path.join(script_dir, 'models', 'det')
        rec_model_dir = os.path.join(script_dir, 'models', 'rec')
        cls_model_dir = os.path.join(script_dir, 'models', 'cls')
        det_model_dir = det_model_dir if os.path.exists(det_model_dir) else None
        rec_model_dir = rec_model_dir if os.path.exists(rec_model_dir) else None
        cls_model_dir = cls_model_dir if os.path.exists(cls_model_dir) else None
        
        # 初始化 PaddleOCR - 优先使用本地模型（只需要 det 和 rec，不需要 cls）
        if det_model_dir and rec_model_dir:
            log_error("检测到本地模型，使用离线模式")
            _paddle_ocr = PaddleOCR(
                lang='ch',
                use_angle_cls=False,
                show_log=False,
                det_model_dir=det_model_dir,
                rec_model_dir=rec_model_dir
            )
        else:
            log_error("未检测到本地模型，首次运行会自动下载（约100MB）")
            _paddle_ocr = PaddleOCR(lang='ch', use_angle_cls=False, show_log=False)
        
        log_error("PaddleOCR初始化完成")
    return _paddle_ocr

def extract_text_from_image(image_path, save_full_result=False):
    try:
        ocr = get_paddle_ocr()
        result = ocr.ocr(image_path)
        full_ocr_result = ""
        
        if result and len(result) > 0 and result[0]:
            all_lines = []
            for idx, line in enumerate(result[0], 1):
                if isinstance(line, (list, tuple)) and len(line) >= 2:
                    if isinstance(line[1], (list, tuple)):
                        text = line[1][0]
                        conf = line[1][1] if len(line[1]) > 1 else 1.0
                    else:
                        text = str(line[1])
                        conf = 1.0
                    all_lines.append(text)
                    if save_full_result:
                        full_ocr_result += f"[{idx}] '{text}' (置信度: {conf:.2f})\n"
            
            if full_ocr_result and save_full_result:
                full_ocr_result = "\n--- OCR完整识别结果 ---\n" + full_ocr_result + "--------------------------\n"
                log_error(full_ocr_result)
            
            return ' '.join(all_lines)
        else:
            if save_full_result:
                log_error("\n--- OCR识别结果: 未检测到文本 ---")
    except Exception as e:
        log_error(f"OCR识别失败: {image_path}, 错误: {e}")
        log_error(traceback.format_exc())
    return ""

def extract_ids_and_names_from_files(files):
    all_ids = []
    all_names = []
    for file_path in files:
        lower_path = file_path.lower()
        if lower_path.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif')):
            ids, text, is_front = extract_id_and_text_from_image(file_path)
            all_ids.extend(ids)
            names = find_names(text)
            all_names.extend(names)
        elif lower_path.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
            ids = find_id_numbers(text, is_id_card_front=False)
            all_ids.extend(ids)
            names = find_names(text)
            all_names.extend(names)
            temp_dir = tempfile.mkdtemp()
            try:
                images = extract_images_from_pdf(file_path, temp_dir)
                for img in images:
                    img_ids, img_text, is_front = extract_id_and_text_from_image(img)
                    all_ids.extend(img_ids)
                    img_names = find_names(img_text)
                    all_names.extend(img_names)
            finally:
                shutil.rmtree(temp_dir, ignore_errors=True)
        elif lower_path.endswith('.docx'):
            text = extract_text_from_docx(file_path)
            ids = find_id_numbers(text, is_id_card_front=False)
            all_ids.extend(ids)
            names = find_names(text)
            all_names.extend(names)
            temp_dir = tempfile.mkdtemp()
            try:
                images = extract_images_from_docx(file_path, temp_dir)
                for img in images:
                    img_ids, img_text, is_front = extract_id_and_text_from_image(img)
                    all_ids.extend(img_ids)
                    img_names = find_names(img_text)
                    all_names.extend(img_names)
            finally:
                shutil.rmtree(temp_dir, ignore_errors=True)
        elif lower_path.endswith('.txt'):
            text = ""
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text += f.read()
            except:
                try:
                    with open(file_path, 'r', encoding='gbk') as f:
                        text += f.read()
                except:
                    pass
            if text:
                ids = find_id_numbers(text, is_id_card_front=False)
                all_ids.extend(ids)
                names = find_names(text)
                all_names.extend(names)
    return list(set(all_ids)), list(set(all_names))

def compare_results(pdf_ids, file_ids, pdf_names, file_names, validate_id):
    """对比PDF和材料中的身份证号码及姓名"""
    matched_ids = []
    unmatched_pdf_ids = []
    unmatched_file_ids = []

    if validate_id:
        def validate_id_number(id_num):
            """
            校验18位身份证号码的校验位
            算法依据ISO 7064:1983.MOD 11-2
            """
            if len(id_num) != 18:
                return False
            # 加权因子
            wi = [7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2]
            # 校验码对应值
            check_code = '10X9876543210'
            # 计算加权和
            total = sum(int(id_num[i]) * wi[i] for i in range(17))
            # 校验校验位
            return check_code[total % 11].upper() == id_num[17].upper()
    else:
        validate_id_number = lambda x: True
    
    valid_pdf_ids = [id for id in pdf_ids if validate_id_number(id)]
    valid_file_ids = [id for id in file_ids if validate_id_number(id)]
    
    matched_ids = [id for id in valid_pdf_ids if id in valid_file_ids]
    unmatched_pdf_ids = [id for id in valid_pdf_ids if id not in valid_file_ids]
    unmatched_file_ids = [id for id in valid_file_ids if id not in valid_pdf_ids]
    
    filtered_pdf_names = filter_hidden_names(pdf_names)
    filtered_file_names = filter_hidden_names(file_names)
    
    matched_names = [name for name in filtered_pdf_names if name in filtered_file_names]
    unmatched_pdf_names = [name for name in filtered_pdf_names if name not in filtered_file_names]
    unmatched_file_names = [name for name in filtered_file_names if name not in filtered_pdf_names]
    
    return matched_ids, unmatched_pdf_ids, unmatched_file_ids, matched_names, unmatched_pdf_names, unmatched_file_names

class WorkerThread(QThread):
    progress = pyqtSignal(int, str)
    single_group_finished = pyqtSignal(int, list, list, list, list, str, str)
    all_finished = pyqtSignal(int, int)
    
    def __init__(self, pdf_files, verify_files, validate_id, check_name):
        super().__init__()
        self.pdf_files = pdf_files
        self.verify_files = verify_files
        self.validate_id = validate_id
        self.check_name = check_name
    
    def run(self):
        global log_content
        log_content = ""
        total_groups = len(self.pdf_files)
        success_count = 0
        failure_count = 0
        try:
            self.progress.emit(0, f"开始处理, 共 {total_groups} 组")
            log_error(f"开始处理, 共 {total_groups} 组")
            
            for i, (pdf_file, verify_file) in enumerate(zip(self.pdf_files, self.verify_files)):
                self.progress.emit(int((i / total_groups) * 20), f"正在处理第 {i+1}/{total_groups} 组...")
                log_error(f"处理第 {i+1} 组: PDF={pdf_file}, 材料={verify_file}")
                
                pdf_ids = []
                pdf_names = []
                file_ids = []
                file_names = []
                failure_reasons = []
                warning_msgs = ""
                
                try:
                    self.progress.emit(int(((i + 0.3) / total_groups) * 100), f"处理第 {i+1}/{total_groups} 组: 提取PDF内容...")
                    log_error("正在提取PDF内容...")
                    
                    pdf_text = extract_text_from_pdf(pdf_file)
                    pdf_ids = find_id_numbers(pdf_text, is_id_card_front=False)
                    if self.check_name:
                        pdf_names = find_names(pdf_text)
                    
                    self.progress.emit(int(((i + 0.5) / total_groups) * 100), f"处理第 {i+1}/{total_groups} 组: 处理核实材料...")
                    log_error("正在处理核实材料...")
                    
                    temp_dir = tempfile.mkdtemp()
                    try:
                        lower_verify = verify_file.lower()
                        all_files = []
                        if lower_verify.endswith(('.zip', '.rar', '.7z')):
                            extracted_dir = process_archive(verify_file, temp_dir)
                            if extracted_dir:
                                all_files = get_all_files(extracted_dir)
                            else:
                                failure_reasons.append("压缩包解压失败，请检查文件格式是否正确或尝试转换为ZIP格式")
                        else:
                            all_files = [verify_file]
                            log_error(f"直接处理文件: {verify_file}")
                        
                        log_error(f"待处理文件列表: {all_files}")
                        if not all_files and not failure_reasons:
                            failure_reasons.append("核实材料解压后未找到任何文件")
                        
                        file_ids, file_names = extract_ids_and_names_from_files(all_files)
                    finally:
                        shutil.rmtree(temp_dir, ignore_errors=True)
                    
                    if not pdf_ids:
                        failure_reasons.append("PDF中未提取到身份证号码, 可能是扫描件或格式不支持")
                    if not file_ids:
                        failure_reasons.append("核实材料中未识别到身份证号码, 请检查图片是否清晰")
                    
                    self.progress.emit(int(((i + 0.9) / total_groups) * 100), f"处理第 {i+1}/{total_groups} 组: 比对结果...")
                    
                    matched_ids, unmatched_pdf_ids, unmatched_file_ids, matched_names, unmatched_pdf_names, unmatched_file_names = compare_results(
                        pdf_ids, file_ids, pdf_names, file_names, self.validate_id
                    )
                    
                    if matched_ids:
                        success_count += 1
                    else:
                        failure_count += 1
                    
                    self.single_group_finished.emit(
                        i + 1,
                        pdf_ids,
                        file_ids,
                        pdf_names,
                        file_names,
                        '; '.join(failure_reasons) if failure_reasons else '无',
                        '; '.join(unmatched_pdf_names) if unmatched_pdf_names else '无'
                    )
                except Exception as e:
                    failure_count += 1
                    log_error(f"处理第 {i+1} 组时发生错误: {e}")
                    self.single_group_finished.emit(
                        i + 1, [], [], [], [],
                        f"处理出错: {str(e)}",
                        '无'
                    )
        except Exception as e:
            log_error(f"处理过程中发生严重错误: {e}")
        finally:
            self.all_finished.emit(success_count, failure_count)

class FileDropListWidget(QListWidget):
    def __init__(self, file_list, parent=None, allowed_extensions=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.file_list = file_list
        self.allowed_extensions = allowed_extensions if allowed_extensions else []
    
    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            event.accept()
    
    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            event.accept()
    
    def dropEvent(self, event):
        added = False
        for url in event.mimeData().urls():
            file_path = url.toLocalFile()
            if not file_path:
                continue
            if not os.path.isfile(file_path):
                continue
            if self.allowed_extensions:
                ext = os.path.splitext(file_path)[1].lower()
                if ext not in self.allowed_extensions:
                    continue
            if file_path not in self.file_list:
                self.file_list.append(file_path)
                self.addItem(file_path)
                added = True
        if added:
            event.acceptProposedAction()
            event.accept()
        else:
            event.ignore()

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.pdf_files = []
        self.verify_files = []
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("身份证号码核对工具 v2.5 (PaddleOCR 2.x)")
        self.setGeometry(100, 100, 900, 700)
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout()
        
        layout.addWidget(QLabel("请求书PDF文件"))
        self.pdf_list = FileDropListWidget(
            file_list=self.pdf_files,
            parent=self,
            allowed_extensions=['.pdf']
        )
        layout.addWidget(self.pdf_list)
        
        btn_layout1 = QHBoxLayout()
        btn_add_pdf = QPushButton("添加PDF文件")
        btn_add_pdf.clicked.connect(self.add_pdf_files)
        btn_layout1.addWidget(btn_add_pdf)
        btn_remove_pdf = QPushButton("删除选中")
        btn_remove_pdf.clicked.connect(lambda: self.remove_selected(self.pdf_list, self.pdf_files))
        btn_layout1.addWidget(btn_remove_pdf)
        layout.addLayout(btn_layout1)
        
        layout.addWidget(QLabel("核实材料(支持图片/压缩包)"))
        self.verify_list = FileDropListWidget(
            file_list=self.verify_files,
            parent=self,
            allowed_extensions=['.zip', '.rar', '.7z', '.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif', '.pdf', '.docx', '.txt']
        )
        layout.addWidget(self.verify_list)
        
        btn_layout2 = QHBoxLayout()
        btn_add_verify = QPushButton("添加核实材料")
        btn_add_verify.clicked.connect(self.add_verify_files)
        btn_layout2.addWidget(btn_add_verify)
        btn_remove_verify = QPushButton("删除选中")
        btn_remove_verify.clicked.connect(lambda: self.remove_selected(self.verify_list, self.verify_files))
        btn_layout2.addWidget(btn_remove_verify)
        layout.addLayout(btn_layout2)
        
        self.validate_checkbox = QCheckBox("校验身份证号码(验证校验位)")
        self.validate_checkbox.setChecked(True)
        layout.addWidget(self.validate_checkbox)
        
        self.name_checkbox = QCheckBox("同时核对姓名")
        self.name_checkbox.setChecked(True)
        layout.addWidget(self.name_checkbox)
        
        btn_layout3 = QHBoxLayout()
        self.start_btn = QPushButton("开始核对")
        self.start_btn.clicked.connect(self.start_processing)
        btn_layout3.addWidget(self.start_btn)
        self.clear_btn = QPushButton("清空全部")
        self.clear_btn.clicked.connect(self.clear_all)
        btn_layout3.addWidget(self.clear_btn)
        layout.addLayout(btn_layout3)
        
        self.progress_bar = QProgressBar()
        layout.addWidget(self.progress_bar)
        
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        layout.addWidget(self.result_text)
        
        central_widget.setLayout(layout)
    
    def add_pdf_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择PDF文件", "", "PDF Files (*.pdf)")
        if files:
            for f in files:
                if f not in self.pdf_files:
                    self.pdf_files.append(f)
                    self.pdf_list.addItem(f)
    
    def add_verify_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "选择核实材料", "",
            "All Files (*.*);;Images (*.png *.jpg *.jpeg *.bmp);;Archives (*.zip *.rar *.7z)"
        )
        if files:
            for f in files:
                if f not in self.verify_files:
                    self.verify_files.append(f)
                    self.verify_list.addItem(f)
    
    def remove_selected(self, list_widget, file_list):
        for item in list_widget.selectedItems():
            row = list_widget.row(item)
            file_list.pop(row)
            list_widget.takeItem(row)
    
    def clear_all(self):
        self.pdf_files.clear()
        self.verify_files.clear()
        self.pdf_list.clear()
        self.verify_list.clear()
        self.result_text.clear()
        self.progress_bar.setValue(0)
    
    def start_processing(self):
        if not self.pdf_files or not self.verify_files:
            QMessageBox.warning(self, "警告", "请先添加PDF文件和核实材料!")
            return
        
        if len(self.pdf_files) != len(self.verify_files):
            QMessageBox.warning(self, "警告", "PDF文件和核实材料数量不匹配!")
            return
        
        self.start_btn.setEnabled(False)
        self.clear_btn.setEnabled(False)
        self.result_text.clear()
        self.progress_bar.setValue(0)
        
        validate_id = self.validate_checkbox.isChecked()
        check_name = self.name_checkbox.isChecked()
        
        self.worker = WorkerThread(
            self.pdf_files.copy(),
            self.verify_files.copy(),
            validate_id,
            check_name
        )
        self.worker.progress.connect(self.update_progress)
        self.worker.single_group_finished.connect(self.on_single_group_finished)
        self.worker.all_finished.connect(self.on_all_finished)
        self.worker.start()
    
    def update_progress(self, value, msg):
        self.progress_bar.setValue(value)
        self.result_text.append(msg)
    
    def on_single_group_finished(self, group_num, pdf_ids, file_ids, pdf_names, file_names, failure_reasons, warning_msgs):
        result = f"\n{'='*60}\n"
        result += f"第 {group_num} 组结果:\n"
        result += f"{'-'*60}\n"
        result += f"【身份证号码】\n"
        result += f"  PDF中: {' | '.join(pdf_ids) if pdf_ids else '无'}\n"
        result += f"  材料中: {' | '.join(file_ids) if file_ids else '无'}\n"
        
        result += f"\n【核对结果】\n"
        if pdf_ids or file_ids:
            for id_num in pdf_ids:
                if id_num in file_ids:
                    result += f"  ✓ 身份证号码匹配成功: {id_num}\n"
            for id_num in pdf_ids:
                if id_num not in file_ids:
                    result += f"  ✗ PDF独有身份证号码: {id_num}\n"
            for id_num in file_ids:
                if id_num not in pdf_ids:
                    result += f"  ✗ 材料独有身份证号码: {id_num}\n"
        else:
            result += f"  ⚠ 未检测到身份证号码\n"
        
        if self.name_checkbox.isChecked():
            hidden_pdf_names = [n for n in pdf_names if is_hidden_name(n)]
            hidden_file_names = [n for n in file_names if is_hidden_name(n)]
            filtered_pdf_names = filter_hidden_names(pdf_names)
            filtered_file_names = filter_hidden_names(file_names)
            
            result += f"\n【姓名】\n"
            result += f"  PDF中: {' | '.join(pdf_names) if pdf_names else '无'}\n"
            if hidden_pdf_names:
                result += f"  (已过滤不公布: {' | '.join(hidden_pdf_names)})\n"
            result += f"  材料中: {' | '.join(file_names) if file_names else '无'}\n"
            if hidden_file_names:
                result += f"  (已过滤不公布: {' | '.join(hidden_file_names)})\n"
            
            if filtered_pdf_names or filtered_file_names:
                for name in filtered_pdf_names:
                    if name in filtered_file_names:
                        result += f"  ✓ 姓名匹配: {name}\n"
                for name in filtered_pdf_names:
                    if name not in filtered_file_names:
                        result += f"  ⚠ PDF独有姓名: {name}\n"
                for name in filtered_file_names:
                    if name not in filtered_pdf_names:
                        result += f"  ⚠ 材料独有姓名: {name}\n"
        
        if failure_reasons and failure_reasons != '无':
            result += f"\n【可能存在的问题】\n"
            for reason in failure_reasons.split('; '):
                if reason != '无':
                    result += f"  ⚠ {reason}\n"
        
        self.result_text.append(result)
    
    def on_all_finished(self, success, failure):
        self.progress_bar.setValue(100)
        summary = f"\n{'='*60}\n"
        summary += f"处理完成! 成功: {success}, 失败: {failure}\n"
        if success > 0 and failure == 0:
            summary += "🎉 所有组的身份证号码均已匹配!\n"
        elif success > 0:
            summary += "⚠ 部分匹配成功, 请检查未匹配的项目。\n"
        else:
            summary += "✗ 未能匹配到任何身份证号码, 请检查材料是否正确。\n"
        summary += f"{'='*60}\n"
        
        self.result_text.append(summary)
        self.start_btn.setEnabled(True)
        self.clear_btn.setEnabled(True)
        
        try:
            with open(error_log_path, "w", encoding="utf-8") as f:
                f.write(log_content)
        except:
            pass

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
