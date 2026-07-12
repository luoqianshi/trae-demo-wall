import os
import re
import random
import zipfile
import shutil
import uuid
import json
import hashlib
import traceback
from datetime import datetime
from functools import wraps
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import RGBColor, Pt
from werkzeug.utils import secure_filename

try:
    import fitz  # PyMuPDF
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    HAS_EXCEL_SUPPORT = True
except ImportError:
    HAS_EXCEL_SUPPORT = False

app = Flask(__name__)
CORS(app)
app.config['SECRET_KEY'] = os.urandom(24)

# 目录
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
OUTPUT_FOLDER = os.path.join(os.path.dirname(__file__), 'outputs')
TEMPLATE_FOLDER = os.path.join(os.path.dirname(__file__), 'templates')
USER_DB_FILE = os.path.join(os.path.dirname(__file__), 'users.json')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(TEMPLATE_FOLDER, exist_ok=True)

PROCESSABLE_EXTENSIONS = {'docx', 'pdf'}
ARCHIVE_EXTENSIONS = {'zip'}

# ────────── 中文字体路径 ──────────
CJK_FONT_PATH = r'C:\Windows\Fonts\msyh.ttc'  # 微软雅黑
CJK_FONT_FALLBACK = r'C:\Windows\Fonts\simhei.ttf'  # 黑体

def get_cjk_font_path():
    if os.path.exists(CJK_FONT_PATH):
        return CJK_FONT_PATH
    if os.path.exists(CJK_FONT_FALLBACK):
        return CJK_FONT_FALLBACK
    return None


# ────────── 姓名提取 ──────────

def extract_name(filename):
    """从文件名中提取学生姓名"""
    name_part = os.path.splitext(filename)[0]

    # 模式1: 数字-中文姓名-其他 (如 1-韩博景-文件-...)
    m = re.search(r'-(?P<name>[\u4e00-\u9fff]{2,4})-', name_part)
    if m:
        return m.group('name')

    # 模式2: 中文姓名_ (如 张三_实验报告)
    m = re.match(r'(?P<name>[\u4e00-\u9fff]{2,4})_', name_part)
    if m:
        return m.group('name')

    # 模式3: 提取最长的连续中文（2字以上）
    chinese_parts = re.findall(r'[\u4e00-\u9fff]+', name_part)
    if chinese_parts:
        best = max(chinese_parts, key=len)
        if len(best) >= 2:
            return best

    # 回退: 去掉常见前缀后缀
    clean = re.sub(r'^(\d+[-_])?\d*[-_]?', '', name_part)
    clean = re.sub(r'[-_](实验报告|文件|作业|报告).*', '', clean)
    clean = re.sub(r'[-_]?\d{6,}[-_]?', '', clean)
    clean = clean.strip('-_ ')
    if clean and len(clean) >= 2:
        return clean

    return name_part[:20]


# ────────── 用户系统 ──────────

def load_users():
    if os.path.exists(USER_DB_FILE):
        with open(USER_DB_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}


def save_users(users):
    with open(USER_DB_FILE, 'w', encoding='utf-8') as f:
        json.dump(users, f, ensure_ascii=False, indent=2)


def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()


def get_current_user():
    token = request.headers.get('X-Auth-Token', '')
    users = load_users()
    for username, data in users.items():
        if data.get('token') == token:
            return username
    return None


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not get_current_user():
            return jsonify({"error": "请先登录"}), 401
        return f(*args, **kwargs)
    return decorated


# ────────── 时间工具 ──────────

def get_formatted_time():
    return datetime.now().strftime("%Y年%m月%d日 %H:%M")


# ────────── Word 文档批注 ──────────

def add_comment_to_docx(doc_path, output_path, comment_text, author="AutoWord",
                        position="end", score=None):
    """在 Word 文档中添加红色批注（支持分数）"""
    doc = Document(doc_path)
    timestamp = get_formatted_time()

    def make_title(para):
        run = para.add_run("【教师批注】")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 0, 0)

    def make_score(para, sc):
        run = para.add_run(f"成绩：{sc} 分")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(255, 0, 0)

    def make_comment(para, text):
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 0, 0)

    def make_footer(para):
        run = para.add_run(
            f"批注人：{author}  |  批注时间：{timestamp}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True

    if position == "start":
        first_para = doc.paragraphs[0] if doc.paragraphs else None

        comment_title = doc.add_paragraph()
        comment_title._element.addprevious(first_para._element if first_para else comment_title._element)
        make_title(comment_title)

        if score is not None:
            score_para = doc.add_paragraph()
            score_para._element.addprevious(comment_title._element)
            make_score(score_para, score)

        comment_para = doc.add_paragraph()
        comment_para._element.addprevious(comment_title._element)
        make_comment(comment_para, comment_text)

        time_para = doc.add_paragraph()
        time_para._element.addprevious(comment_para._element)
        make_footer(time_para)
    else:
        doc.add_paragraph()
        comment_title = doc.add_paragraph()
        make_title(comment_title)

        if score is not None:
            doc.add_paragraph()
            score_para = doc.add_paragraph()
            make_score(score_para, score)

        comment_para = doc.add_paragraph()
        make_comment(comment_para, comment_text)

        doc.add_paragraph()
        time_para = doc.add_paragraph()
        make_footer(time_para)

    doc.save(output_path)
    return True


# ────────── PDF 文档批注 ──────────

def add_comment_to_pdf(pdf_path, output_path, comment_text, author="AutoWord",
                       position="end", score=None):
    """在 PDF 中添加红色批注（支持中文）"""
    if not HAS_PDF_SUPPORT:
        raise RuntimeError("缺少 PyMuPDF 库，无法处理 PDF 文件")

    doc = fitz.open(pdf_path)
    timestamp = get_formatted_time()

    if position == "start":
        new_page = doc.new_page(0, width=595, height=842)
    else:
        new_page = doc.new_page(-1, width=595, height=842)

    # 注册中文字体
    font_path = get_cjk_font_path()
    CJK_FONT_NAME = "custom_cjk"
    if font_path:
        try:
            new_page.insert_font(fontname=CJK_FONT_NAME, fontfile=font_path)
            font_args = {"fontname": CJK_FONT_NAME}
        except Exception:
            font_args = {}
    else:
        font_args = {}

    # 标题
    title_rect = fitz.Rect(60, 60, 535, 100)
    new_page.insert_textbox(
        title_rect, "【教师批注】", fontsize=14,
        color=(1, 0, 0), **font_args
    )

    # 分数（如有）
    y_offset = 110
    if score is not None:
        score_rect = fitz.Rect(60, 110, 535, 145)
        new_page.insert_textbox(
            score_rect, f"成绩：{score} 分", fontsize=13,
            color=(1, 0, 0), **font_args
        )
        y_offset = 155

    # 批注内容
    body_rect = fitz.Rect(60, y_offset, 535, 650)
    new_page.insert_textbox(
        body_rect, comment_text, fontsize=11,
        color=(1, 0, 0), **font_args
    )

    # 页脚
    footer_rect = fitz.Rect(60, 700, 535, 760)
    new_page.insert_textbox(
        footer_rect,
        f"批注人：{author}  |  批注时间：{timestamp}",
        fontsize=9, color=(0.5, 0.5, 0.5), **font_args
    )

    doc.save(output_path)
    doc.close()
    return True


# ────────── 文件调度（带成功/失败分离） ──────────

def process_single_file(file_path, success_dir, failed_dir,
                        comment_text, author="AutoWord", position="end", score=None):
    """处理单个文件，成功放入 success_dir，失败复制原文件到 failed_dir 并返回错误信息"""
    filename = os.path.basename(file_path)
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    student_name = extract_name(filename)

    try:
        # 确保输出目录存在
        os.makedirs(success_dir, exist_ok=True)
        os.makedirs(failed_dir, exist_ok=True)

        success_path = os.path.join(success_dir, filename)

        if ext == 'docx':
            add_comment_to_docx(file_path, success_path, comment_text, author, position, score)
        elif ext == 'pdf':
            add_comment_to_pdf(file_path, success_path, comment_text, author, position, score)
        return {"filename": filename, "status": "success", "name": student_name, "score": score}
    except Exception as e:
        # 复制原文件到失败目录
        try:
            os.makedirs(failed_dir, exist_ok=True)
            failed_path = os.path.join(failed_dir, filename)
            shutil.copy2(file_path, failed_path)
        except Exception:
            pass
        error_msg = str(e)
        # 简化常见错误信息
        if "Bad CRC-32" in error_msg or "File is not a zip file" in error_msg:
            error_msg = "文件已损坏，无法打开"
        elif "Package not found" in error_msg:
            error_msg = "Word 文档格式异常或已损坏"
        elif "cannot open document" in error_msg.lower():
            error_msg = "PDF 文档已损坏或格式不支持"
        return {"filename": filename, "status": "error", "error": error_msg, "name": student_name}


# ────────── 嵌套压缩包递归扫描 ──────────

def safe_extract_zip(zip_path, extract_dir):
    """安全解压 ZIP，自动处理 GBK/UTF-8 编码的中文文件名"""
    with zipfile.ZipFile(zip_path, 'r') as zf:
        for info in zf.infolist():
            filename = info.filename

            # 检测是否为 GBK 编码的文件名
            # ZIP 规范使用 CP437 编码文件名，中文 Windows 常使用 GBK
            try:
                raw_bytes = filename.encode('cp437')
                raw_bytes.decode('utf-8')
                # UTF-8 解码成功，无需修改
            except (UnicodeDecodeError, UnicodeEncodeError):
                # 尝试 GBK / GB18030 解码
                try:
                    raw = filename.encode('cp437')
                    info.filename = raw.decode('gbk')
                except (UnicodeDecodeError, UnicodeEncodeError):
                    try:
                        raw = filename.encode('cp437')
                        info.filename = raw.decode('gb18030')
                    except (UnicodeDecodeError, UnicodeEncodeError):
                        pass  # 保持原名

        os.makedirs(extract_dir, exist_ok=True)

        # 逐个提取，使用修正后的文件名
        for info in zf.infolist():
            extracted_path = os.path.join(extract_dir, info.filename)
            # 安全检查：防止路径遍历
            if not os.path.abspath(extracted_path).startswith(os.path.abspath(extract_dir)):
                continue
            os.makedirs(os.path.dirname(extracted_path), exist_ok=True)
            with zf.open(info.orig_filename) as src, open(extracted_path, 'wb') as dst:
                dst.write(src.read())


def recursively_extract_and_collect(extract_root, collected):
    """递归扫描目录，展开所有嵌套压缩包"""
    MAX_NESTING = 10
    for _ in range(MAX_NESTING):
        found_zip = False
        for root, dirs, files in os.walk(extract_root):
            for file in files:
                if file.lower().endswith('.zip'):
                    zip_path = os.path.join(root, file)
                    nested_dir = os.path.join(root, file[:-4])
                    try:
                        safe_extract_zip(zip_path, nested_dir)
                        os.remove(zip_path)
                        found_zip = True
                    except Exception:
                        pass
        if not found_zip:
            break

    for root, dirs, files in os.walk(extract_root):
        for file in files:
            ext = file.rsplit('.', 1)[1].lower() if '.' in file else ''
            if ext in PROCESSABLE_EXTENSIONS and not file.startswith('~$'):
                collected.append(os.path.join(root, file))


def process_zip_archive(zip_path, output_dir, comment_text, author="AutoWord",
                        position="end", score_min=None, score_max=None):
    """处理 ZIP 压缩包，成功和失败文件分别存放，支持分数分配"""
    extract_dir = os.path.join(output_dir, '_extracted')
    os.makedirs(extract_dir, exist_ok=True)

    results = []

    try:
        safe_extract_zip(zip_path, extract_dir)
    except zipfile.BadZipFile:
        return [{"filename": os.path.basename(zip_path),
                 "status": "error",
                 "error": "损坏的压缩包，无法解压"}]

    collected_files = []
    recursively_extract_and_collect(extract_dir, collected_files)

    # 成功和失败的输出目录
    success_dir = os.path.join(output_dir, 'success')
    failed_dir = os.path.join(output_dir, 'failed')

    # 为每个文件分配分数
    has_scores = score_min is not None and score_max is not None
    file_scores = {}
    if has_scores:
        for file_path in collected_files:
            file_scores[file_path] = random.randint(int(score_min), int(score_max))

    for file_path in collected_files:
        rel_path = os.path.relpath(file_path, extract_dir)
        rel_dir = os.path.dirname(rel_path)

        file_success_dir = os.path.join(success_dir, rel_dir) if rel_dir else success_dir
        file_failed_dir = os.path.join(failed_dir, rel_dir) if rel_dir else failed_dir

        score = file_scores.get(file_path) if has_scores else None
        result = process_single_file(file_path, file_success_dir, file_failed_dir,
                                     comment_text, author, position, score)
        results.append(result)

    shutil.rmtree(extract_dir, ignore_errors=True)
    return results


def create_output_zip(source_dir, zip_path):
    """将目录中的文件打包成 ZIP（使用 UTF-8 编码文件名）"""
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(source_dir):
            for file in files:
                if not file.endswith('.zip'):
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, source_dir)
                    zipf.write(file_path, arcname)
        # 确认文件名编码为 UTF-8
        return zipf


def create_excel_report(results, output_path):
    """生成成绩统计 Excel 表"""
    if not HAS_EXCEL_SUPPORT:
        return None

    # 收集有分数的成功结果
    scored = [r for r in results if r.get('status') == 'success' and r.get('score') is not None]
    if not scored:
        return None

    wb = Workbook()
    ws = wb.active
    ws.title = "成绩统计"

    # 样式定义
    header_font = Font(bold=True, color="FFFFFF", size=12)
    header_fill = PatternFill("solid", fgColor="2563EB")
    header_align = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style="thin", color="D9DEE7"),
        right=Side(style="thin", color="D9DEE7"),
        top=Side(style="thin", color="D9DEE7"),
        bottom=Side(style="thin", color="D9DEE7"),
    )
    zebra_fill_1 = PatternFill("solid", fgColor="FFFFFF")
    zebra_fill_2 = PatternFill("solid", fgColor="F7F9FC")
    center_align = Alignment(horizontal="center", vertical="center")

    # 表头
    headers = ["序号", "姓名", "文件名", "成绩"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    # 按姓名排序
    scored.sort(key=lambda x: x.get('name', '') or x['filename'])

    # 数据行
    for row_idx, r in enumerate(scored, 2):
        fill = zebra_fill_1 if (row_idx - 2) % 2 == 0 else zebra_fill_2
        name = r.get('name', '')
        filename = r['filename']
        score = r['score']

        cells = [
            (1, row_idx - 1),  # 序号
            (2, name),         # 姓名
            (3, filename),     # 文件名
            (4, score),        # 成绩
        ]
        for col, value in cells:
            cell = ws.cell(row=row_idx, column=col, value=value)
            cell.border = thin_border
            cell.fill = fill
            cell.alignment = center_align if col in (1, 4) else Alignment(horizontal="left", vertical="center")

    # 统计行
    stat_row = len(scored) + 2
    ws.cell(row=stat_row, column=1, value="统计")
    ws.cell(row=stat_row, column=1).font = Font(bold=True)
    ws.cell(row=stat_row, column=2, value=f"共 {len(scored)} 人")
    ws.cell(row=stat_row, column=2).font = Font(bold=True)

    scores = [r['score'] for r in scored]
    ws.cell(row=stat_row, column=3, value=f"平均分: {sum(scores)/len(scores):.1f}")
    ws.cell(row=stat_row, column=3).font = Font(bold=True)
    ws.cell(row=stat_row, column=4, value=f"最高分: {max(scores)} / 最低分: {min(scores)}")
    ws.cell(row=stat_row, column=4).font = Font(bold=True)

    # 列宽
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 16
    ws.column_dimensions['C'].width = 40
    ws.column_dimensions['D'].width = 10

    wb.save(output_path)
    return output_path


def allowed_file(filename):
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    return ext in PROCESSABLE_EXTENSIONS or ext in ARCHIVE_EXTENSIONS


# ────────── 预设模板 ──────────

DEFAULT_TEMPLATES = [
    {
        "id": "default", "name": "通用批注", "category": "通用",
        "content": "实验报告已批阅。整体完成度良好，请根据课堂讲评内容进一步完善实验分析部分。如有疑问，欢迎课后交流。"
    },
    {
        "id": "good", "name": "优秀评价", "category": "评价",
        "content": "实验报告完成质量优秀！实验步骤清晰，数据分析到位，结论合理。继续保持！"
    },
    {
        "id": "needs_improvement", "name": "需改进", "category": "评价",
        "content": "实验报告已完成，但存在以下问题需要改进：\n1. 实验步骤描述不够详细\n2. 数据分析部分需要补充\n3. 结论部分需要与实验数据更紧密结合\n请修改后重新提交。"
    },
    {
        "id": "excellent", "name": "特别优秀", "category": "评价",
        "content": "非常出色的实验报告！不仅完成了所有要求，还展现了深入的思考和创新能力。实验设计合理，数据记录完整，分析透彻。这份报告可以作为范本供其他同学参考。"
    },
    {
        "id": "late_submission", "name": "迟交提醒", "category": "提醒",
        "content": "实验报告已收到，但提交时间已超过截止日期。请注意下次按时提交，以免影响平时成绩。如有特殊情况，请提前与任课教师沟通。"
    },
    {
        "id": "format_issue", "name": "格式问题", "category": "提醒",
        "content": "实验报告内容尚可，但格式不符合要求。请注意：\n1. 标题格式应统一\n2. 图表需添加编号和说明\n3. 参考文献格式需规范\n请按模板要求修改后重新提交。"
    },
    {
        "id": "encouragement", "name": "鼓励", "category": "鼓励",
        "content": "看得出你在实验中付出了努力，虽然还有一些不足，但整体方向是正确的。建议多参考教材中的范例，加强与理论知识的结合。相信下次会有更好的表现，加油！"
    },
    {
        "id": "plagiarism_warning", "name": "抄袭警告", "category": "警告",
        "content": "经检测，本报告与其他同学的报告高度相似，涉嫌抄袭。学术诚信是基本要求，请独立完成实验报告。请于3日内重新提交原创报告，否则将按相关规定处理。"
    },
]


def ensure_default_templates():
    for t in DEFAULT_TEMPLATES:
        path = os.path.join(TEMPLATE_FOLDER, f"{t['id']}.txt")
        if not os.path.exists(path):
            with open(path, 'w', encoding='utf-8') as f:
                f.write(t['content'])
        meta_path = os.path.join(TEMPLATE_FOLDER, f"{t['id']}.json")
        if not os.path.exists(meta_path):
            with open(meta_path, 'w', encoding='utf-8') as f:
                json.dump({"name": t['name'], "category": t['category']}, f, ensure_ascii=False)


ensure_default_templates()


# ────────── API 路由 ──────────

@app.route('/api/auth/register', methods=['POST'])
def register():
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    if not username or not password:
        return jsonify({"error": "用户名和密码不能为空"}), 400
    users = load_users()
    if username in users:
        return jsonify({"error": "用户名已存在"}), 409
    users[username] = {
        "password_hash": hash_password(password),
        "token": None,
        "created_at": get_formatted_time()
    }
    save_users(users)
    return jsonify({"success": True, "message": "注册成功"})


@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    users = load_users()
    if username not in users:
        return jsonify({"error": "用户名或密码错误"}), 401
    if users[username]["password_hash"] != hash_password(password):
        return jsonify({"error": "用户名或密码错误"}), 401
    token = hashlib.sha256(os.urandom(32)).hexdigest()
    users[username]["token"] = token
    save_users(users)
    return jsonify({"success": True, "token": token, "username": username})


@app.route('/api/auth/logout', methods=['POST'])
@login_required
def logout():
    username = get_current_user()
    users = load_users()
    if username in users:
        users[username]["token"] = None
        save_users(users)
    return jsonify({"success": True})


@app.route('/api/auth/me', methods=['GET'])
def get_me():
    username = get_current_user()
    if not username:
        return jsonify({"logged_in": False})
    return jsonify({"logged_in": True, "username": username})


@app.route('/api/templates', methods=['GET'])
def get_templates():
    templates = []
    if os.path.exists(TEMPLATE_FOLDER):
        for f in os.listdir(TEMPLATE_FOLDER):
            if f.endswith('.txt'):
                template_id = f[:-4]
                with open(os.path.join(TEMPLATE_FOLDER, f), 'r', encoding='utf-8') as tf:
                    content = tf.read()
                name = template_id
                category = "自定义"
                meta_path = os.path.join(TEMPLATE_FOLDER, f"{template_id}.json")
                if os.path.exists(meta_path):
                    with open(meta_path, 'r', encoding='utf-8') as mf:
                        meta = json.load(mf)
                        name = meta.get('name', template_id)
                        category = meta.get('category', '自定义')
                templates.append({"id": template_id, "name": name,
                                  "category": category, "content": content})
    templates.sort(key=lambda x: (x['category'], x['name']))
    return jsonify({"templates": templates})


@app.route('/api/templates', methods=['POST'])
@login_required
def save_template():
    data = request.json
    template_name = data.get('name', 'custom').strip()
    template_content = data.get('content', '').strip()
    template_category = data.get('category', '自定义').strip()
    if not template_name or not template_content:
        return jsonify({"error": "模板名称和内容不能为空"}), 400
    template_id = template_name.replace(' ', '_').replace('/', '_')
    with open(os.path.join(TEMPLATE_FOLDER, f"{template_id}.txt"), 'w', encoding='utf-8') as f:
        f.write(template_content)
    with open(os.path.join(TEMPLATE_FOLDER, f"{template_id}.json"), 'w', encoding='utf-8') as f:
        json.dump({"name": template_name, "category": template_category}, f, ensure_ascii=False)
    return jsonify({"success": True, "id": template_id})


@app.route('/api/templates/<template_id>', methods=['DELETE'])
@login_required
def delete_template(template_id):
    preset_ids = [t['id'] for t in DEFAULT_TEMPLATES]
    if template_id in preset_ids:
        return jsonify({"error": "预设模板不能删除"}), 403
    txt_path = os.path.join(TEMPLATE_FOLDER, f"{template_id}.txt")
    meta_path = os.path.join(TEMPLATE_FOLDER, f"{template_id}.json")
    if os.path.exists(txt_path):
        os.remove(txt_path)
    if os.path.exists(meta_path):
        os.remove(meta_path)
    return jsonify({"success": True})


@app.route('/api/process', methods=['POST'])
def process_files():
    """处理上传的文件，返回成功和失败两个下载链接"""
    if 'file' not in request.files:
        return jsonify({"error": "没有上传文件"}), 400

    file = request.files['file']
    comment_text = request.form.get('comment', '已批阅')
    author = request.form.get('author', 'AutoWord')
    position = request.form.get('position', 'end')

    if file.filename == '':
        return jsonify({"error": "文件名为空"}), 400
    if not allowed_file(file.filename):
        return jsonify({"error": "仅支持 .docx / .pdf 文件或 .zip 压缩包"}), 400

    task_id = str(uuid.uuid4())[:8]
    task_upload_dir = os.path.join(UPLOAD_FOLDER, task_id)
    task_output_dir = os.path.join(OUTPUT_FOLDER, task_id)
    os.makedirs(task_upload_dir, exist_ok=True)
    os.makedirs(task_output_dir, exist_ok=True)

    raw_filename = file.filename
    filename = secure_filename(raw_filename)
    # secure_filename strips CJK chars, resulting in just the extension or empty
    # If the result is just an extension or empty, use original with path chars replaced
    if not filename or ('.' in raw_filename and filename.lower() == raw_filename.rsplit('.', 1)[1].lower()):
        filename = raw_filename.replace('/', '_').replace('\\', '_').replace(':', '_')
    if not filename:
        return jsonify({"error": "文件名无效"}), 400
    file_path = os.path.join(task_upload_dir, filename)
    file.save(file_path)

    # 分数范围
    score_min = request.form.get('score_min')
    score_max = request.form.get('score_max')
    if score_min:
        try:
            score_min = int(score_min)
        except ValueError:
            score_min = None
    if score_max:
        try:
            score_max = int(score_max)
        except ValueError:
            score_max = None

    file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

    if file_ext == 'zip':
        results = process_zip_archive(
            file_path, task_output_dir, comment_text, author, position,
            score_min, score_max
        )
    else:
        success_dir = os.path.join(task_output_dir, 'success')
        failed_dir = os.path.join(task_output_dir, 'failed')
        score = random.randint(score_min, score_max) if score_min is not None and score_max is not None else None
        result = process_single_file(
            file_path, success_dir, failed_dir,
            comment_text, author, position, score
        )
        results = [result]

    # 分别打包成功和失败文件
    success_dir = os.path.join(task_output_dir, 'success')
    failed_dir = os.path.join(task_output_dir, 'failed')

    success_zip_path = os.path.join(OUTPUT_FOLDER, f"{task_id}_success.zip")
    failed_zip_path = os.path.join(OUTPUT_FOLDER, f"{task_id}_failed.zip")

    has_success = os.path.exists(success_dir) and any(os.scandir(success_dir))
    has_failed = os.path.exists(failed_dir) and any(os.scandir(failed_dir))

    if has_success:
        create_output_zip(success_dir, success_zip_path)
    if has_failed:
        create_output_zip(failed_dir, failed_zip_path)
    elif os.path.exists(failed_zip_path):
        os.remove(failed_zip_path)

    success_count = sum(1 for r in results if r['status'] == 'success')
    error_count = len(results) - success_count

    # 分离成功和失败的详细结果
    success_results = [r for r in results if r['status'] == 'success']
    failed_results = [r for r in results if r['status'] == 'error']

    # 生成 Excel 成绩表
    excel_path = os.path.join(OUTPUT_FOLDER, f"{task_id}_scores.xlsx")
    excel_created = False
    if HAS_EXCEL_SUPPORT and success_results:
        excel_created = create_excel_report(results, excel_path) is not None
    if not excel_created and os.path.exists(excel_path):
        os.remove(excel_path)

    return jsonify({
        "task_id": task_id,
        "success": True,
        "total": len(results),
        "success_count": success_count,
        "error_count": error_count,
        "success_results": success_results,
        "failed_results": failed_results,
        "download_url": f"/download/{task_id}/success" if has_success else None,
        "failed_download_url": f"/download/{task_id}/failed" if has_failed else None,
        "excel_url": f"/download/{task_id}/excel" if excel_created else None,
    })


@app.route('/api/download/<task_id>/<result_type>', methods=['GET'])
def download_result(task_id, result_type):
    """下载成功文件、失败文件或成绩表"""
    if result_type == 'success':
        zip_path = os.path.join(OUTPUT_FOLDER, f"{task_id}_success.zip")
        download_name = "autoword_批注成功.zip"
        mime = 'application/zip'
    elif result_type == 'failed':
        zip_path = os.path.join(OUTPUT_FOLDER, f"{task_id}_failed.zip")
        download_name = "autoword_批注失败（原文件）.zip"
        mime = 'application/zip'
    elif result_type == 'excel':
        zip_path = os.path.join(OUTPUT_FOLDER, f"{task_id}_scores.xlsx")
        download_name = "autoword_成绩统计表.xlsx"
        mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    else:
        return jsonify({"error": "无效的下载类型"}), 400

    if not os.path.exists(zip_path):
        return jsonify({"error": "文件不存在或已过期"}), 404
    return send_file(zip_path, as_attachment=True, download_name=download_name, mimetype=mime)


@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "ok",
        "pdf_support": HAS_PDF_SUPPORT,
        "excel_support": HAS_EXCEL_SUPPORT,
        "version": "3.2"
    })


# ────────── 前端页面服务 ──────────

FRONTEND_DIR = os.path.join(os.path.dirname(__file__), '..', 'frontend')

@app.route('/')
def serve_index():
    return send_file(os.path.join(FRONTEND_DIR, 'index.html'))


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
