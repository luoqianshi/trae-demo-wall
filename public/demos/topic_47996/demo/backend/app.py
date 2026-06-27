from flask import Flask, request, jsonify, session
from flask_cors import CORS
import sqlite3
import hashlib
import os
from datetime import datetime
import tempfile
import json
import random
import requests
# from faker import Faker
# fake = Faker('zh_CN')  # 使用中文本地化
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'  # 在生产环境中应该使用环境变量
CORS(app, supports_credentials=True)

# 火山引擎方舟API配置
ARK_API_KEY = "ark-343755d6-3b45-4f13-86b0-b10570063c3d-d0540"
ARK_API_BASE = "https://ark.cn-beijing.volces.com/api/v3/chat/completions"
ARK_MODEL = "ep-20260627125005-bzkf9"

def call_ark_api(messages, temperature=0.2, max_tokens=500, timeout=120):
    """调用火山引擎方舟API"""
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {ARK_API_KEY}"
    }
    payload = {
        "model": ARK_MODEL,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    response = requests.post(ARK_API_BASE, headers=headers, json=payload, timeout=timeout)
    response.raise_for_status()
    result = response.json()
    return result["choices"][0]["message"]["content"].strip()

# 初始化Faker生成中文姓名
# fake = Faker('zh_CN')

# 数据库初始化 - 使用try.py的方法
def init_database():
    """初始化数据库和表结构"""
    conn = sqlite3.connect('medical_records.db')
    c = conn.cursor()
    
    # 创建用户表
    c.execute('''CREATE TABLE IF NOT EXISTS yiliao_user (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
             )''')
    
    # 创建医疗记录表 - 添加user_id字段（修复表结构问题）
    c.execute('''
        CREATE TABLE IF NOT EXISTS yiliao_jilu (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            role TEXT,
            content TEXT,
            mode TEXT,
            score REAL,
            FOREIGN KEY(user_id) REFERENCES yiliao_user(id)
        )
    ''')
    
    # 创建病历表 - 添加user_id字段
    c.execute('''
        CREATE TABLE IF NOT EXISTS bingli (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            patient_info TEXT,
            consultation_history TEXT,
            FOREIGN KEY(user_id) REFERENCES yiliao_user(id)
        )
    ''')
    
    conn.commit()
    conn.close()

# 获取当前用户ID
def get_current_user_id(username):
    """获取当前用户ID"""
    try:
        conn = sqlite3.connect('medical_records.db')
        c = conn.cursor()
        c.execute("SELECT id FROM yiliao_user WHERE username = ?", (username,))
        user = c.fetchone()
        return user[0] if user else None
    except Exception as e:
        print(f"获取用户ID错误: {str(e)}")
        return None
    finally:
        conn.close()

# 哈希密码函数
def hash_password(password):
    """使用SHA256哈希密码"""
    return hashlib.sha256(password.encode()).hexdigest()

# 使用TF-IDF作为嵌入替代方案
def generate_tfidf_embeddings(texts):
    """使用TF-IDF生成文本嵌入（替代方案）"""
    if not hasattr(generate_tfidf_embeddings, 'vectorizer'):
        generate_tfidf_embeddings.vectorizer = TfidfVectorizer(max_features=2000)
        if isinstance(texts, str):
            texts = [texts]
        generate_tfidf_embeddings.vectorizer.fit(texts)
    
    if isinstance(texts, str):
        texts = [texts]
    
    embeddings = generate_tfidf_embeddings.vectorizer.transform(texts).toarray()
    return np.array(embeddings)

# 向量存储类
class VectorStore:
    def __init__(self):
        self.documents = []
        self.embeddings = None
    
    def add_texts(self, texts):
        """添加文本并生成嵌入"""
        self.documents.extend(texts)
        new_embeddings = generate_tfidf_embeddings(texts)
        
        if self.embeddings is None:
            self.embeddings = new_embeddings
        else:
            self.embeddings = np.vstack([self.embeddings, new_embeddings])
    
    def similarity_search(self, query, k=3):
        """搜索相似文档"""
        query_embedding = generate_tfidf_embeddings([query])
        similarities = cosine_similarity(query_embedding, self.embeddings)[0]
        indices = np.argsort(similarities)[::-1][:k]
        return [self.documents[i] for i in indices]

# 知识库管理
class KnowledgeBase:
    def __init__(self):
        self.vector_store = VectorStore()
        self.load_documents()
    
    def load_documents(self):
        """加载知识库文档"""
        medical_knowledge = [
            "胰腺癌（pancreatic cancer）是胰腺上皮来源的恶性肿瘤，最常见为胰腺导管腺癌（PDAC），占全部病例的85-90%。",
            "胰腺癌分子分型包括：经典型、准间质型、纯基底样型和鳞状样型。不同分子亚型对治疗反应存在差异。",
            "胰腺癌的典型症状包括：无痛性黄疸（库尔伏尔氏定律）、上腹部或腰背部持续性疼痛（夜间加重）、体重减轻（>10%）、食欲减退、脂肪泻、新发糖尿病或原有糖尿病恶化。",
            "胰头癌患者以黄疸（80-90%）和瘙痒为主要表现；胰体尾癌患者多以腹痛和体重减轻为主要表现。",
            "胰腺癌诊断标准（2023 NCCN指南）：影像学发现胰腺占位+CA19-9显著升高（排除胆道梗阻）+病理学确认（尽可能获取组织学诊断）。",
            "病理诊断要求：活检应包含足够的组织量，需报告组织学类型、分化程度、微卫星不稳定状态（MSI）、肿瘤突变负荷（TMB）等分子标志物。",
            "AJCC第8版TNM分期系统适用：T1（≤2cm）、T2（>2且≤4cm）、T3（>4cm）、T4（侵犯腹腔干或肠系膜上动脉）",
            "术后主要预后因素：R0切除、淋巴结阴性（N0）、CA19-9术后降至正常范围、分化良好、新辅助化疗后病理缓解。",
            "根治性手术：胰头癌-Whipple手术（PD）；胰体尾癌-远端胰腺切除+脾切除；全胰腺切除适应证严格（双侧病变或胰腺实质广泛受累）",
            "新辅助治疗：适用于临界可切除/Borderline Resectable胰腺癌（BRPC），方案：FOLFIRINOX（首选）或吉西他滨+白蛋白紫杉醇+放疗",
            "辅助治疗：所有R0切除患者均应接受辅助化疗，方案：改良FOLFIRINOX（体能状态良好者）或吉西他滨+卡培他滨（CA19-9>1000U/mL）",
            "晚期一线治疗：吉西他滨+白蛋白结合型紫杉醇或FOLFIRINOX（PS 0-1分）；存在BRCA1/2突变者考虑奥拉帕利维持治疗",
            "术后并发症标准化处理：ISGPS定义的术后胰瘘（POPF）分三级管理，B/C级需CT引导下引流+抗生素",
            "术后随访：第1年每3个月，2-3年每6个月，4-5年每年1次，包含体检、CA19-9、增强CT/MRI",
            "免疫治疗：针对微卫星不稳定（MSI-H）胰腺癌，PD-1抑制剂（帕博利珠单抗）获得FDA突破性疗法认定",
            "靶向治疗：KRAS G12C抑制剂索托拉西布+EGFR抑制剂联合方案在KRAS G12C突变型胰腺癌中显示42%疾病控制率"
        ]
        self.vector_store.add_texts(medical_knowledge)
    
    def query_knowledge(self, question):
        """查询知识库获取相关信息"""
        docs = self.vector_store.similarity_search(question, k=3)
        context = "\n".join(docs)
        return context

# 练习系统类
class PracticeSystem:
    def __init__(self, knowledge_base):
        self.knowledge_base = knowledge_base
        # 基础评分标准
        self.scoring_criteria = {
            "greeting": {"权重": 10, "子项": {"礼貌称呼": 2, "自我介绍": 2, "问候患者": 3, "建立信任": 3}},
            "symptom_inquiry": {"权重": 30, "子项": {"症状持续时间": 4, "症状性质": 4, "症状程度": 4, "症状变化": 4, "伴随症状": 4, "诱发缓解因素": 4, "与进食关系": 3, "与体位关系": 3}},
            "history_inquiry": {"权重": 15, "子项": {"既往疾病史": 3, "手术史": 3, "家族史": 2, "过敏史": 2, "药物史": 2, "生活史": 3}},
            "diagnosis_accuracy": {"权重": 20, "子项": {"鉴别诊断": 5, "体征检查": 5, "辅助检查建议": 5, "初步诊断": 5}},
            "treatment_plan": {"权重": 25, "子项": {"治疗解释": 2, "手术建议": 5, "药物建议": 5, "随访计划": 5, "生活建议": 5}},
            "critical_miss": {"权重": 0, "扣分项": {"未问黄疸": -3, "未问体重变化": -2, "未问疼痛性质": -2, "未问发热情况": -2, "未问大小便情况": -1, "未问既往病史": -3, "未问家族史": -2, "未问过敏史": -1, "未做体格检查": -5, "诊断错误": -10, "治疗方案不当": -8, "辱骂患者": -15, "令患者愤怒": -8, "令患者沮丧": -5, "态度恶劣": -10, "不当言论": -6}},
            "advanced_diagnosis": {"权重": 15, "子项": {"正确识别可切除性": 3, "建议恰当影像检查": 3, "鉴别诊断全面性": 3, "肿瘤标志物应用": 3, "分期准确性": 3}}
        }
        
        # 病例特异性评分标准
        self.case_specific_criteria = {
            "A1_腹泻": {
                "关键问诊点": {
                    "工作压力与症状关系": 5,
                    "大便性状描述": 4,
                    "发作频率和持续时间": 4,
                    "伴随症状(脐周痛、肠鸣)": 3,
                    "既往检查结果": 3,
                    "睡眠和焦虑状况": 3
                },
                "人文关怀要点": {
                    "回应癌症担忧": 5,
                    "解释肠镜必要性": 4,
                    "缓解检查恐惧": 4
                }
            },
            "A2_咯血": {
                "关键问诊点": {
                    "痰血性状和诱因": 5,
                    "胸痛特点": 4,
                    "吸烟史详细询问": 5,
                    "家族肺癌史": 4,
                    "既往治疗效果": 3
                },
                "人文关怀要点": {
                    "戒烟建议和指导": 5,
                    "肺癌筛查解释": 4,
                    "家族史相关安慰": 4
                }
            },
            "A3_尿路感染": {
                "关键问诊点": {
                    "尿路症状三联征": 5,
                    "诱发因素(劳累熬夜)": 4,
                    "月经史和性生活史": 4,
                    "既往类似发作": 3,
                    "用药史和效果": 3
                },
                "人文关怀要点": {
                    "解释感染原因": 4,
                    "生活指导建议": 4,
                    "检查必要性说明": 3
                }
            },
            "A4_胸痛": {
                "关键问诊点": {
                    "胸痛诱发因素变化": 5,
                    "疼痛性质和放射": 5,
                    "硝酸甘油效果变化": 5,
                    "高胆固醇血症史": 4,
                    "家族心血管病史": 3
                },
                "人文关怀要点": {
                    "心梗风险解释": 5,
                    "紧急检查安排": 4,
                    "治疗方案讨论": 4
                }
            },
            "B1_腰痛": {
                "关键问诊点": {
                    "疼痛性质和诱因": 4,
                    "体位关系": 4,
                    "既往影像学检查": 4,
                    "职业相关因素": 3,
                    "治疗效果评估": 3
                },
                "人文关怀要点": {
                    "疾病预后解释": 4,
                    "生活方式指导": 4,
                    "治疗方案制定": 3
                }
            },
            "B2_腹痛": {
                "关键问诊点": {
                    "疼痛性质和部位": 5,
                    "月经史详细询问": 5,
                    "阴道流血情况": 5,
                    "妊娠史和避孕史": 4,
                    "伴随症状": 4
                },
                "人文关怀要点": {
                    "宫外孕风险解释": 5,
                    "紧急处理安排": 5,
                    "心理支持": 4
                }
            },
            "B3_右上腹痛": {
                "关键问诊点": {
                    "疼痛诱因(油腻饮食)": 5,
                    "疼痛性质和放射": 4,
                    "发热和寒战": 4,
                    "既往检查结果": 4,
                    "家族胆石症史": 3
                },
                "人文关怀要点": {
                    "手术必要性解释": 4,
                    "避免盲目止痛": 5,
                    "进一步检查安排": 4
                }
            },
            "B4_发热咳嗽": {
                "关键问诊点": {
                    "发热特点和用药": 5,
                    "咳嗽性质和痰液": 4,
                    "喂养和发育史": 4,
                    "接触史和诱因": 3,
                    "伴随症状": 3
                },
                "人文关怀要点": {
                    "发热与智力关系解释": 5,
                    "合理用药指导": 4,
                    "家长焦虑缓解": 4
                }
            }
        }
        
        self.core_questions = {
            "黄疸": ["皮肤发黄", "眼睛发黄", "小便黄"],
            "体重变化": ["瘦了多少", "体重减轻", "体重下降"],
            "疼痛": ["哪里痛", "怎么痛", "夜间痛", "后背痛"],
            "消化": ["不想吃", "吃不下", "拉肚子", "消化不良"],
            "病史": ["抽烟", "喝酒", "家里人癌症", "糖尿病"]
        }
        
        self.detailed_scores = {
            "greeting": {"礼貌称呼": 0.0, "自我介绍": 0.0, "问候患者": 0.0, "建立信任": 0.0},
            "symptom_inquiry": {"症状持续时间": 0.0, "症状性质": 0.0, "症状程度": 0.0, "症状变化": 0.0, "伴随症状": 0.0, "诱发缓解因素": 0.0, "与进食关系": 0.0, "与体位关系": 0.0},
            "history_inquiry": {"既往疾病史": 0.0, "手术史": 0.0, "家族史": 0.0, "过敏史": 0.0, "药物史": 0.0, "生活史": 0.0},
            "diagnosis_accuracy": {"鉴别诊断": 0.0, "体征检查": 0.0, "辅助检查建议": 0.0, "初步诊断": 0.0},
            "treatment_plan": {"治疗解释": 0.0, "手术建议": 0.0, "药物建议": 0.0, "随访计划": 0.0, "生活建议": 0.0},
            "advanced_diagnosis": {"正确识别可切除性": 0.0, "建议恰当影像检查": 0.0, "鉴别诊断全面性": 0.0, "肿瘤标志物应用": 0.0, "分期准确性": 0.0}
        }
        
        self.current_score = 0.00
        self.score_history = []
        self.feedback = []
        
        self.emotion_state = {
            "emotion": "平静",
            "intensity": 0.5,
            "money_concern": 0.5,
            "memory": []
        }
        
        self.patient_profile = self.generate_new_patient_profile()
        self.patient_personality = self.generate_patient_personality()
    
    def generate_patient_personality(self):
        """生成详细患者人格特征"""
        personalities = [
            {"description": "友善耐心的中年患者", "emotion": "平静", "dialect": "普通话", "provoke_chance": 0.1, "money_concern": 0.3, "traits": ["合作", "信任医生"]},
            {"description": "焦虑易怒的退休老人", "emotion": "烦躁", "dialect": "东北话", "provoke_chance": 0.6, "money_concern": 0.8, "traits": ["节俭", "疑病症"]},
            {"description": "粗鲁挑剔的老板", "emotion": "愤怒", "dialect": "四川话", "provoke_chance": 0.8, "money_concern": 0.4, "traits": ["自我中心", "要求高"]},
            {"description": "地方口音重的农民", "emotion": "困惑", "dialect": "闽南话", "provoke_chance": 0.3, "money_concern": 0.9, "traits": ["经济困难", "医疗知识缺乏"]},
            {"description": "疑心重的知识分子", "emotion": "怀疑", "dialect": "湖南话", "provoke_chance": 0.7, "money_concern": 0.5, "traits": ["查过资料", "要求解释"]},
            {"description": "抑郁担忧的家庭照顾者", "emotion": "悲伤", "dialect": "河南话", "provoke_chance": 0.4, "money_concern": 0.7, "traits": ["担心家庭", "害怕花钱"]},
            

            {"description": "乐观的年轻上班族", "emotion": "积极", "dialect": "标准普通话", "provoke_chance": 0.1, "money_concern": 0.6, "traits": ["有医保", "相信现代医学"]},
            {"description": "固执的老年人", "emotion": "倔强", "dialect": "山东话", "provoke_chance": 0.9, "money_concern": 0.8, "traits": ["拒绝检查", "相信偏方"]},
            {"description": "精打细算的小店主", "emotion": "谨慎", "dialect": "浙江话", "provoke_chance": 0.3, "money_concern": 0.95, "traits": ["询问每项费用", "讨价还价"]},
            {"description": "过度紧张的教师", "emotion": "神经质", "dialect": "标准普通话", "provoke_chance": 0.5, "money_concern": 0.4, "traits": ["过度解读症状", "频繁确认"]},
            {"description": "沉默寡言的工人", "emotion": "内向", "dialect": "河北话", "provoke_chance": 0.2, "money_concern": 0.85, "traits": ["惜字如金", "害怕医疗账单"]},
            {"description": "多疑的公务员", "emotion": "猜疑", "dialect": "北京话", "provoke_chance": 0.6, "money_concern": 0.5, "traits": ["质疑诊断", "索要检查报告"]},
            {"description": "恐慌的单亲家长", "emotion": "恐惧", "dialect": "湖北话", "provoke_chance": 0.2, "money_concern": 0.9, "traits": ["担心孩子", "极端节俭"]},
            {"description": "傲慢的企业高管", "emotion": "轻视", "dialect": "标准普通话", "provoke_chance": 0.8, "money_concern": 0.2, "traits": ["不遵医嘱", "强调地位"]},
            {"description": "迷信的农村居民", "emotion": "神秘", "dialect": "江西话", "provoke_chance": 0.3, "money_concern": 0.9, "traits": ["相信鬼神", "抗拒现代医疗"]},
            {"description": "虚弱的老人", "emotion": "疲惫", "dialect": "四川话", "provoke_chance": 0.1, "money_concern": 0.9, "traits": ["动作缓慢", "反复确认费用"]},
            
            {"description": "健谈的出租车司机", "emotion": "健谈", "dialect": "天津话", "provoke_chance": 0.4, "money_concern": 0.7, "traits": ["分享经历", "抱怨物价"]},
            {"description": "悲观的失业者", "emotion": "绝望", "dialect": "重庆话", "provoke_chance": 0.3, "money_concern": 0.95, "traits": ["放弃治疗倾向", "多次询问费用"]},
            {"description": "愤世嫉俗的艺术家", "emotion": "嘲讽", "dialect": "台湾腔", "provoke_chance": 0.7, "money_concern": 0.5, "traits": ["质疑系统", "拖延缴费"]},
            {"description": "虔诚的宗教信徒", "emotion": "平和", "dialect": "云南话", "provoke_chance": 0.1, "money_concern": 0.6, "traits": ["经常祈祷", "接受命运"]},
            {"description": "控制狂的企业主", "emotion": "强势", "dialect": "广东话", "provoke_chance": 0.8, "money_concern": 0.3, "traits": ["要求掌控一切", "质疑用药"]},
            {"description": "天真的学生", "emotion": "好奇", "dialect": "标准普通话", "provoke_chance": 0.1, "money_concern": 0.8, "traits": ["多问题", "家庭资助"]},

            {"description": "吝啬的会计师", "emotion": "计较", "dialect": "上海话", "provoke_chance": 0.4, "money_concern": 0.98, "traits": ["精确计算费用", "要求明细"]},
            {"description": "顺从的家庭成员", "emotion": "被动", "dialect": "江苏话", "provoke_chance": 0.1, "money_concern": 0.85, "traits": ["依赖配偶决定", "心疼花钱"]},
            {"description": "好斗的退伍军人", "emotion": "对抗", "dialect": "东北话", "provoke_chance": 0.9, "money_concern": 0.4, "traits": ["挑战权威", "忽略费用"]},
            {"description": "焦虑的准父母", "emotion": "紧张", "dialect": "湖南话", "provoke_chance": 0.3, "money_concern": 0.6, "traits": ["担忧孩子", "咨询补助"]},
            
            {"description": "世故的商人", "emotion": "圆滑", "dialect": "福建话", "provoke_chance": 0.2, "money_concern": 0.5, "traits": ["试图拉关系", "要求折扣"]},
            {"description": "无力的残障人士", "emotion": "无助", "dialect": "河北话", "provoke_chance": 0.1, "money_concern": 0.92, "traits": ["依赖救助", "担忧负担"]},
            {"description": "过度礼貌的外地人", "emotion": "谨慎", "dialect": "广西话", "provoke_chance": 0.05, "money_concern": 0.9, "traits": ["不断道歉", "敏感价格"]},
            {"description": "骄傲的前运动员", "emotion": "否认", "dialect": "标准普通话", "provoke_chance": 0.7, "money_concern": 0.3, "traits": ["忽视症状", "拒绝承认严重性"]},
            {"description": "恐慌的心脏病患者", "emotion": "恐惧", "dialect": "陕西话", "provoke_chance": 0.4, "money_concern": 0.75, "traits": ["频繁恐慌发作", "担心费用"]},
            {"description": "乐观的癌症幸存者", "emotion": "希望", "dialect": "普通话", "provoke_chance": 0.05, "money_concern": 0.6, "traits": ["积极态度", "感激医保"]},
            {"description": "固执的糖尿病患者", "emotion": "顽固", "dialect": "闽南话", "provoke_chance": 0.85, "money_concern": 0.7, "traits": ["饮食不规律", "抱怨药价"]},
            
            {"description": "不信任的医疗受害者", "emotion": "敌意", "dialect": "四川话", "provoke_chance": 0.9, "money_concern": 0.6, "traits": ["质疑能力", "索要费用说明"]},
            {"description": "困惑的老年痴呆患者", "emotion": "混乱", "dialect": "河南话", "provoke_chance": 0.3, "money_concern": 0.8, "traits": ["重复问题", "依赖家属"]},
            {"description": "霸道的家长", "emotion": "专横", "dialect": "东北话", "provoke_chance": 0.9, "money_concern": 0.4, "traits": ["要求特殊照顾", "不关注价格"]},
            {"description": "受宠的独生子女", "emotion": "任性", "dialect": "标准普通话", "provoke_chance": 0.7, "money_concern": 0.1, "traits": ["依赖父母支付", "要求快速见效"]},
            {"description": "节俭的退休教师", "emotion": "理性", "dialect": "山东话", "provoke_chance": 0.1, "money_concern": 0.9, "traits": ["比较治疗方案", "选择性检查"]},
            {"description": "害怕的外籍工人", "emotion": "胆怯", "dialect": "湖南话", "provoke_chance": 0.1, "money_concern": 0.95, "traits": ["语言障碍", "担忧账单"]},
            {"description": "过度关注的健康追求者", "emotion": "偏执", "dialect": "标准普通话", "provoke_chance": 0.6, "money_concern": 0.5, "traits": ["要求额外检查", "忽视费用"]},
            
            {"description": "沮丧的长期病患者", "emotion": "冷漠", "dialect": "江西话", "provoke_chance": 0.2, "money_concern": 0.85, "traits": ["治疗疲惫", "抱怨花费"]},
            {"description": "独立的自由职业者", "emotion": "防备", "dialect": "广东话", "provoke_chance": 0.4, "money_concern": 0.7, "traits": ["自费医疗", "谨慎开支"]},
            {"description": "悲观的临终患者", "emotion": "绝望", "dialect": "河北话", "provoke_chance": 0.2, "money_concern": 0.9, "traits": ["拒绝治疗", "考虑花费"]},
            {"description": "疑病的年轻上班族", "emotion": "焦虑", "dialect": "上海话", "provoke_chance": 0.5, "money_concern": 0.6, "traits": ["上网自诊", "担心昂贵检查"]},
            {"description": "固执的胃炎患者", "emotion": "顽固", "dialect": "闽南话", "provoke_chance": 0.8, "money_concern": 0.5, "traits": ["不忌口", "抱怨药价"]},
            {"description": "害怕手术的公务员", "emotion": "恐惧", "dialect": "标准普通话", "provoke_chance": 0.3, "money_concern": 0.4, "traits": ["询问替代方案", "担忧请假"]},
            {"description": "强势的知识分子", "emotion": "挑剔", "dialect": "台湾腔", "provoke_chance": 0.85, "money_concern": 0.3, "traits": ["查阅最新研究", "要求定制方案"]},
            
            {"description": "顺从的老年患者", "emotion": "温顺", "dialect": "湖北话", "provoke_chance": 0.05, "money_concern": 0.9, "traits": ["完全信任", "心疼花费"]},
            {"description": "多疑的营养师", "emotion": "批判", "dialect": "标准普通话", "provoke_chance": 0.7, "money_concern": 0.6, "traits": ["质疑饮食建议", "关注药价"]},
            {"description": "焦虑的老烟民", "emotion": "懊悔", "dialect": "山西话", "provoke_chance": 0.4, "money_concern": 0.7, "traits": ["难以戒烟", "担心医药费"]},
            {"description": "过度保护的父母", "emotion": "紧张", "dialect": "陕西话", "provoke_chance": 0.3, "money_concern": 0.5, "traits": ["为孩子担忧", "不惜花费"]},
            {"description": "愤怒的医疗纠纷者", "emotion": "暴怒", "dialect": "四川话", "provoke_chance": 0.95, "money_concern": 0.9, "traits": ["既往不满", "威胁诉讼"]},
            {"description": "不信任的民间医者", "emotion": "傲慢", "dialect": "湖南话", "provoke_chance": 0.9, "money_concern": 0.4, "traits": ["轻视西医", "偏好偏方"]},
            {"description": "抑郁的离婚人士", "emotion": "消沉", "dialect": "河南话", "provoke_chance": 0.2, "money_concern": 0.85, "traits": ["无病感", "财务担忧"]},
            
            {"description": "固执的乙肝患者", "emotion": "防御", "dialect": "东北话", "provoke_chance": 0.75, "money_concern": 0.6, "traits": ["隐瞒病史", "治疗不规律"]},
            {"description": "害怕麻醉的老年", "emotion": "恐惧", "dialect": "山东话", "provoke_chance": 0.3, "money_concern": 0.7, "traits": ["拒绝手术", "多次确认风险"]},

            {"description": "精明的保险用户", "emotion": "计算", "dialect": "江苏话", "provoke_chance": 0.1, "money_concern": 0.95, "traits": ["了解保险条款", "要求医保内项目"]},
            {"description": "急躁的餐饮业主", "emotion": "不耐烦", "dialect": "粤语", "provoke_chance": 0.8, "money_concern": 0.6, "traits": ["赶时间", "抱怨等候"]},
            {"description": "隐忍的慢性痛患者", "emotion": "坚忍", "dialect": "江西话", "provoke_chance": 0.1, "money_concern": 0.9, "traits": ["长期忍受", "仅基础治疗"]},
            {"description": "恐惧化疗的癌症患者", "emotion": "恐惧", "dialect": "标准普通话", "provoke_chance": 0.4, "money_concern": 0.7, "traits": ["犹豫不决", "担心副作用花费"]},
            {"description": "强势的法律人士", "emotion": "威胁", "dialect": "北京话", "provoke_chance": 0.85, "money_concern": 0.3, "traits": ["提法律条款", "要求签字证明"]},
            {"description": "谨慎的退休教师", "emotion": "谨慎", "dialect": "广东话", "provoke_chance": 0.2, "money_concern": 0.7, "traits": ["注重专业知识", "重视患者"]},
            {"description": "敏感的情感患者", "emotion": "敏感", "dialect": "湖南话", "provoke_chance": 0.6, "money_concern": 0.8, "traits": ["关注隐私", "不分享信息"]},
            {"description": "强势的法律人士", "emotion": "威胁", "dialect": "北京话", "provoke_chance": 0.85, "money_concern": 0.3, "traits": ["提法律条款", "要求签字证明"]}
        ]
        personality = random.choice(personalities)
        self.emotion_state["emotion"] = personality["emotion"]
        self.emotion_state["money_concern"] = personality["money_concern"]
        return personality
    
    def generate_new_patient_profile(self):
        """生成新的患者信息 - 随机组合症状与病程"""
        case_type = random.choice(['A1_腹泻', 'A2_咯血', 'A3_尿路感染', 'A4_胸痛',
                                   'B1_腰痛', 'B2_腹痛', 'B3_右上腹痛', 'B4_发热咳嗽'])
        return self._build_randomized_case(case_type)
    
    def _build_randomized_case(self, case_type):
        """根据病例类型随机生成症状组合与病程"""
        name = random.choice(['张伟', '李娜', '王强', '赵敏', '陈明', '刘芳', '周杰', '吴丽', '黄勇', '林秀'])
        age_config = {
            'A1_腹泻': (25, 55), 'A2_咯血': (35, 70), 'A3_尿路感染': (20, 50),
            'A4_胸痛': (40, 75), 'B1_腰痛': (30, 65), 'B2_腹痛': (18, 40),
            'B3_右上腹痛': (30, 65), 'B4_发热咳嗽': (1, 12)
        }
        age = random.randint(*age_config.get(case_type, (20, 60)))
        # B2_腹痛 为宫外孕病例，强制女性
        gender = '女' if case_type == 'B2_腹痛' else random.choice(['男', '女'])
        occupations = ['教师', '职员', '工程师', '销售', '退休', '学生', '自由职业', '司机', '厨师', '会计']
        occupation = random.choice(occupations)
        
        # 症状池
        symptom_pools = self._get_symptom_pools()
        pool = symptom_pools[case_type]
        
        main_symptoms = random.sample(pool['main_options'], min(3, len(pool['main_options'])))
        accompany_symptoms = random.sample(pool['accompany_options'], min(3, len(pool['accompany_options'])))
        systemic_symptoms = random.sample(pool['systemic_options'], min(2, len(pool['systemic_options'])))
        
        course = self._generate_course(case_type, pool)
        
        profile = {
            'case_type': case_type,
            'basic_info': {'姓名': name, '年龄': age, '性别': gender, '职业': occupation},
            'symptoms': {
                '主要症状': main_symptoms,
                '伴随症状': accompany_symptoms,
                '全身症状': systemic_symptoms
            },
            'course': course
        }
        
        profile['lab_results'] = self.generate_lab_results(case_type)
        return profile
    
    def _get_symptom_pools(self):
        """各病例类型的症状池"""
        return {
            'A1_腹泻': {
                'main_options': [
                    '间断腹泻8个月', '反复腹泻6个月', '慢性腹泻1年', '间歇性腹泻10个月',
                    '腹泻伴腹痛5个月', '反复稀便7个月', '饭后腹泻9个月', '晨起腹泻4个月'
                ],
                'accompany_options': [
                    '脐周隐痛', '肠鸣亢进', '含黏液便', '未消化食物残渣',
                    '腹部胀气', '排便不尽感', '里急后重', '下腹坠胀'
                ],
                'systemic_options': [
                    '焦虑不安', '失眠多梦', '工作压力大', '食欲减退',
                    '体重下降', '精神紧张', '疲劳乏力', '注意力不集中'
                ],
                'course_templates': [
                    '腹泻{months}个月，{trigger}时易发作，近{recent}加重',
                    '反复腹泻{months}个月，{trigger}后症状明显，每日{times}次',
                    '{months}个月前开始腹泻，{trigger}时加重，伴有{extra}',
                    '间歇性腹泻{months}个月，近{recent}{trigger}后发作频繁'
                ],
                'triggers': ['工作压力大', '情绪紧张', '饮食不当', '劳累', '熬夜', '考试前'],
                'extras': ['腹痛', '肠鸣', '黏液便', '体重下降']
            },
            'A2_咯血': {
                'main_options': [
                    '痰中带血1个月', '咯血2周', '反复痰血3个月', '咳血丝痰6周',
                    '间断咯血2个月', '晨起痰中带血1个月', '咯鲜血2周', '痰血交替3个月'
                ],
                'accompany_options': [
                    '右下胸隐痛', '咳嗽加重', '胸闷不适', '劳累后加重',
                    '左侧胸痛', '深呼吸时胸痛', '气短', '咳白色黏痰'
                ],
                'systemic_options': [
                    '精神稍差', '吸烟20年', '体重下降', '低热',
                    '夜间盗汗', '食欲减退', '疲劳乏力', '声音嘶哑'
                ],
                'course_templates': [
                    '反复痰中带血{months}个月，近{recent}出现{extra}',
                    '{months}个月前开始咯血，{trigger}后加重',
                    '痰血交替出现{months}个月，近{recent}{extra}',
                    '间断咯血{months}个月，{trigger}诱发，伴有{extra}'
                ],
                'triggers': ['劳累', '感冒', '气候变化', '吸烟', '熬夜'],
                'extras': ['右下胸隐痛', '胸闷', '咳嗽加重', '气短']
            },
            'A3_尿路感染': {
                'main_options': [
                    '尿频尿急尿痛1天', '排尿灼痛2天', '尿频伴烧灼感1天',
                    '尿急尿痛半天', '尿频尿不尽感2天', '排尿困难伴刺痛1天'
                ],
                'accompany_options': [
                    '每小时排尿1次', '每次量少', '尿道烧灼感', '下腹不适',
                    '终末血尿', '排尿不尽感', '耻骨上压痛', '尿液浑浊'
                ],
                'systemic_options': [
                    '疲劳乏力', '睡眠差', '近期装修劳累', '饮水不足',
                    '近期熬夜', '工作繁忙', '饮食不规律', '抵抗力下降'
                ],
                'course_templates': [
                    '突发{main}，{trigger}',
                    '{main}，{trigger}后出现，近{recent}加重',
                    '突然出现{main}，{trigger}',
                    '{main}，{trigger}后发作，排尿时{extra}'
                ],
                'triggers': ['近期装修劳累', '饮水不足', '熬夜后', '着凉后', '经期后', '劳累后'],
                'extras': ['烧灼感明显', '疼痛难忍', '无法正常工作', '夜尿增多']
            },
            'A4_胸痛': {
                'main_options': [
                    '间断胸痛3个月', '反复胸痛6个月', '阵发胸痛2个月',
                    '胸骨后压榨痛1个月', '胸闷胸痛5个月', '活动后胸痛4个月'
                ],
                'accompany_options': [
                    '向左肩放射', '寒冷诱发', '情绪激动诱发', '劳累后加重',
                    '休息后缓解', '心悸', '压迫感', '烧心感'
                ],
                'systemic_options': [
                    '精神紧张', '高胆固醇血症5年', '高血压病史', '睡眠差',
                    '糖尿病史', '吸烟史', '家族冠心病史', '肥胖'
                ],
                'course_templates': [
                    '间断胸痛{months}个月，{trigger}诱发，近{recent}加重',
                    '反复胸痛{months}个月，{trigger}时发作，{extra}',
                    '{months}个月前开始胸痛，{trigger}后加重，{extra}',
                    '阵发性胸痛{months}个月，近{recent}{trigger}后发作频繁'
                ],
                'triggers': ['寒冷', '情绪激动', '劳累', '饱餐后', '爬楼梯', '快走'],
                'extras': ['向左肩放射', '休息后缓解', '含服硝酸甘油缓解', '持续时间延长']
            },
            'B1_腰痛': {
                'main_options': [
                    '反复腰痛2年', '慢性腰痛3年', '腰骶部疼痛1年半',
                    '间歇性腰痛5年', '持续性腰痛8个月', '劳累后腰痛2年'
                ],
                'accompany_options': [
                    '久坐加重', '平卧缓解', '活动受限', '下肢放射痛',
                    '弯腰困难', '翻身疼痛', '行走不便', '臀部酸痛'
                ],
                'systemic_options': [
                    '睡眠差', '长期久坐', '缺乏运动', '肥胖',
                    '体力劳动', '精神压力大', '姿势不良', '空调房工作'
                ],
                'course_templates': [
                    '腰痛{months}个月，{trigger}加重，近{recent}{extra}',
                    '反复腰痛{months}个月，{trigger}，{extra}',
                    '{months}个月来腰痛反复发作，{trigger}时加重',
                    '腰痛病史{months}个月，近{recent}{trigger}后{extra}'
                ],
                'triggers': ['久坐', '搬重物', '弯腰', '受凉', '劳累', '晨起'],
                'extras': ['活动受限', '下肢放射痛', '行走困难', '腰部僵硬']
            },
            'B2_腹痛': {
                'main_options': [
                    '下腹痛5小时', '左下腹痛8小时', '突发下腹痛6小时',
                    '右下腹痛3小时', '腹部剧痛4小时', '下腹撕裂痛7小时'
                ],
                'accompany_options': [
                    '恶心呕吐', '肛门坠胀感', '阴道少量流血', '头晕心慌',
                    '面色苍白', '出冷汗', '血压下降', '腹部压痛'
                ],
                'systemic_options': [
                    '末次月经45天前', '月经推迟2周', '妊娠可能', '月经不规律',
                    '近期性生活', '避孕失败', '停经6周', '早孕反应'
                ],
                'course_templates': [
                    '突发{main}，{trigger}',
                    '{main}，{trigger}后出现，{extra}',
                    '突然{main}，{trigger}，{extra}',
                    '{main}，{trigger}，伴有{extra}'
                ],
                'triggers': ['无明显诱因', '剧烈运动后', '性生活后', '劳累后', '排便后'],
                'extras': ['恶心呕吐', '肛门坠胀', '头晕', '出冷汗']
            },
            'B3_右上腹痛': {
                'main_options': [
                    '右上腹痛半天', '右上腹绞痛1天', '右上腹阵痛6小时',
                    '右上腹胀痛8小时', '右上腹持续性痛2天', '右上腹剧痛10小时'
                ],
                'accompany_options': [
                    '向右肩胛放射', '恶心呕吐', '油腻饮食后发作', '发热',
                    '腹胀', '嗳气', '反酸', '黄疸'
                ],
                'systemic_options': [
                    '自觉发热', '精神萎靡', '食欲不振', '睡眠差',
                    '既往胆结石史', '体重偏胖', '高脂饮食', '胆囊炎病史'
                ],
                'course_templates': [
                    '{trigger}后出现{main}，{extra}',
                    '{main}，{trigger}诱发，{extra}',
                    '突然{main}，{trigger}，{extra}',
                    '{main}，{trigger}后发作，{extra}'
                ],
                'triggers': ['油腻饮食', '饱餐', '饮酒', '进食油炸食品', '吃火锅', '吃烧烤'],
                'extras': ['向右肩放射', '恶心', '呕吐', '发热']
            },
            'B4_发热咳嗽': {
                'main_options': [
                    '发热3天咳嗽2天', '发热2天咳嗽3天', '发热1天咳嗽伴痰鸣',
                    '反复发热5天咳嗽4天', '发热2天干咳1天', '发热伴咳嗽4天'
                ],
                'accompany_options': [
                    '39.7℃高热', '喉中痰鸣', '38.5℃发热', '干咳无痰',
                    '咳黄色黏痰', '鼻塞', '咽痛', '呼吸急促'
                ],
                'systemic_options': [
                    '流涕喷嚏', '干呕', '进食减少', '烦躁哭闹',
                    '精神萎靡', '食欲减退', '睡眠不安', '呕吐'
                ],
                'course_templates': [
                    '{trigger}后受凉，{main}',
                    '{trigger}，{main}，{extra}',
                    '{trigger}后出现{main}，{extra}',
                    '{main}，{trigger}后加重，{extra}'
                ],
                'triggers': ['公园游玩', '幼儿园传染', '天气变化', '游泳', '外出淋雨', '吹空调'],
                'extras': ['喉中痰鸣', '流涕', '进食减少', '精神差']
            }
        }
    
    def _generate_course(self, case_type, pool):
        """根据病例类型随机生成病程描述"""
        import re
        template = random.choice(pool['course_templates'])
        trigger = random.choice(pool['triggers'])
        extra = random.choice(pool['extras'])
        
        # 从模板中提取需要的参数
        result = template
        
        if '{months}' in result:
            months = random.choice([3, 4, 5, 6, 7, 8, 9, 10, 12])
            result = result.replace('{months}', str(months))
        if '{main}' in result:
            result = result.replace('{main}', random.choice(pool['main_options']))
        if '{trigger}' in result:
            result = result.replace('{trigger}', trigger)
        if '{extra}' in result:
            result = result.replace('{extra}', extra)
        if '{recent}' in result:
            recent = random.choice(['1周', '3天', '2周', '1个月', '数日'])
            result = result.replace('{recent}', recent)
        if '{times}' in result:
            result = result.replace('{times}', random.choice(['3-4次', '4-5次', '5-6次', '2-3次']))
        
        return result
    
    def generate_lab_results(self, case_type):
        """根据病例类型生成相应的检查结果"""
        lab_results_templates = {
            'A1_腹泻': [
                "粪常规：正常",
                "血常规：正常",
                "建议：肠镜检查排除器质性疾病"
            ],
            'A2_咯血': [
                "胸片：待查",
                "血常规：正常",
                "建议：胸部CT检查，戒烟"
            ],
            'A3_尿路感染': [
                "尿常规：白细胞+++",
                "尿培养：待查",
                "建议：多饮水，抗感染治疗"
            ],
            'A4_胸痛': [
                "心电图：ST段改变",
                "心肌酶：轻度升高",
                "建议：紧急心血管评估"
            ],
            'B1_腰痛': [
                "腰椎MRI：椎间盘突出",
                "X线：腰椎退行性改变",
                "建议：理疗，避免久坐"
            ],
            'B2_腹痛': [
                "血HCG：阳性",
                "B超：宫外孕可能",
                "建议：紧急妇科会诊"
            ],
            'B3_右上腹痛': [
                "血常规：白细胞升高",
                "B超：胆囊结石，胆囊肿大",
                "建议：外科会诊"
            ],
            'B4_发热咳嗽': [
                "血常规：白细胞正常",
                "胸片：肺纹理增粗",
                "建议：对症治疗，观察"
            ]
        }
        
        return lab_results_templates.get(case_type, ["检查结果待完善"])
    
    def update_emotion(self, doctor_input):
        """根据医生的问题更新患者情感状态"""
        # 处理特殊情况（医生不当行为）
        if doctor_input in ["doctor_abuse", "doctor_rude", "doctor_dismissive"]:
            if doctor_input == "doctor_abuse":
                self.emotion_state["emotion"] = "暴怒"
                self.emotion_state["intensity"] = 1.0
                self.emotion_state["memory"].append("医生使用了辱骂性语言，患者非常愤怒")
            elif doctor_input == "doctor_rude":
                self.emotion_state["emotion"] = "愤怒"
                self.emotion_state["intensity"] = 0.8
                self.emotion_state["memory"].append("医生态度恶劣，患者感到愤怒")
            elif doctor_input == "doctor_dismissive":
                self.emotion_state["emotion"] = "沮丧"
                self.emotion_state["intensity"] = 0.7
                self.emotion_state["memory"].append("医生轻视症状，患者感到沮丧")
            return
        
        positive_words = ["感谢", "放心", "安心", "帮助", "关心", "安慰", "理解"]
        negative_words = ["严重", "癌症", "手术", "费用", "花钱", "风险", "死亡"]
        provocative_words = ["错误", "不可能", "不懂", "胡说", "没听清"]
        
        positive_impact = sum(1 for word in positive_words if word in doctor_input) * 0.1
        negative_impact = sum(1 for word in negative_words if word in doctor_input) * 0.2
        provocative_impact = sum(1 for word in provocative_words if word in doctor_input) * 0.3
        
        net_impact = positive_impact - (negative_impact + provocative_impact)
        self.emotion_state["intensity"] = max(0, min(1, self.emotion_state["intensity"] + net_impact))
        
        if self.emotion_state["intensity"] > 0.7:
            if self.patient_personality["emotion"] == "平静":
                self.emotion_state["emotion"] = "满意"
            elif self.patient_personality["emotion"] in ["烦躁", "愤怒"]:
                self.emotion_state["emotion"] = "愤怒"
        elif self.emotion_state["intensity"] < 0.3:
            if self.patient_personality["emotion"] == "困惑":
                self.emotion_state["emotion"] = "焦虑"
            elif self.patient_personality["emotion"] == "悲伤":
                self.emotion_state["emotion"] = "抑郁"
    
    def generate_patient_response(self, doctor_input):
        """根据医生提问生成患者回答"""
        self.update_emotion(doctor_input)
        self.emotion_state["memory"].append(f"医生: {doctor_input}")
        if len(self.emotion_state["memory"]) > 5:
            self.emotion_state["memory"] = self.emotion_state["memory"][-5:]
        
        basic = self.patient_profile['basic_info']
        symptoms = self.patient_profile['symptoms']
        personality = self.patient_personality
        emotion = self.emotion_state
        
        money_prompt = ""
        if emotion["money_concern"] > 0.5 and random.random() < 0.3:
            money_prompt = "特别注意：你对医疗费用非常敏感，会询问每个检查的费用。"
        
        symptom_state = self.get_symptom_state(len(emotion["memory"]))
        
        context = f"""
        你是一名胰腺癌患者，正在接受医生问诊。
        当前情绪：{emotion['emotion']}（强度{emotion['intensity']:.1f}) {money_prompt}
        患者人格：{personality['description']}，口音：{personality['dialect']}
        基本资料：{basic['姓名']}，{basic['年龄']}岁，{basic['性别']}性
        主要症状：{', '.join(symptoms['主要症状'])}
        伴随症状：{', '.join(symptoms['伴随症状'])}
        全身症状：{', '.join(symptoms['全身症状'])}
        病程：{self.patient_profile['course']}
        人格特质：{', '.join(personality['traits'])}
        
        [病情阶段：{symptom_state["phase"]}]
        [疾病进展：{symptom_state["description"]}]
        [当前症状变化：{symptom_state["changes"]}]
        
        最近对话:
        {chr(10).join(emotion["memory"][-3:])}
        
        请根据医生的问题，以患者身份回答。要求：
        1. 符合你的人格特征和情绪状态
        2. 逐步透露症状信息，不要一次性说完
        3. 用患者的语言风格，带有地方口音特色
        4. 回答长度控制在50-100字
        5. 可以适当表现出对费用的担心
        6. 根据情绪状态调整语气（平静/焦虑/愤怒等）
        """
        
        try:
            response = call_ark_api(
                messages=[
                    {"role": "system", "content": context},
                    {"role": "user", "content": doctor_input},
                ],
                temperature=0.9,
                max_tokens=200,
            )
            return response
        except Exception as e:
            print(f"API调用失败: {str(e)}")
            expressions = [
                "（擦擦额头上的汗）", "（痛苦地皱眉）", "（捂着肚子）", 
                "（叹气）", "（摇头）", "（咳嗽几声）", "（强忍疼痛）"
            ]
            responses = [
                f"俺肚子疼得厉害{random.choice(expressions)}",
                f"你个大夫会不会看病啊！{random.choice(expressions)}",
                f"最近老是恶心不想吃饭{random.choice(expressions)}",
                f"这啥时候才能好啊，愁死人了{random.choice(expressions)}",
                f"俺都瘦了十多斤了，浑身没力气{random.choice(expressions)}"
            ]
            return random.choice(responses)
    
    def get_symptom_state(self, dialogue_round):
        """根据对话轮次动态调整症状描述"""
        states = [
            {"phase": "初始阶段", "description": "轻度症状", "changes": "仅描述基本症状"},
            {"phase": "发展阶段", "description": "中度症状进展", "changes": "症状加重（3周）"},
            {"phase": "关键阶段", "description": "出现并发症", "changes": "新症状（黄疸加重/疼痛增加）"},
            {"phase": "诊断阶段", "description": "需要诊断决策", "changes": "等待医生诊断"},
        ]
        return states[min(dialogue_round//3, len(states)-1)]
    
    def get_current_case_type(self):
        """根据患者症状识别当前病例类型（优先使用存储的case_type）"""
        # 优先使用 profile 中存储的 case_type
        if 'case_type' in self.patient_profile:
            return self.patient_profile['case_type']
        
        # 回退到关键词匹配
        symptoms = self.patient_profile.get('symptoms', {})
        main_symptoms = symptoms.get('主要症状', [])
        age = self.patient_profile.get('basic_info', {}).get('年龄', 999)
        
        if any('腹泻' in symptom for symptom in main_symptoms):
            return 'A1_腹泻'
        elif any('咯血' in symptom or '痰中带血' in symptom for symptom in main_symptoms):
            return 'A2_咯血'
        elif any('尿频' in symptom or '尿急' in symptom or '尿痛' in symptom for symptom in main_symptoms):
            return 'A3_尿路感染'
        elif any('胸痛' in symptom for symptom in main_symptoms):
            return 'A4_胸痛'
        elif any('腰痛' in symptom for symptom in main_symptoms):
            return 'B1_腰痛'
        elif any('下腹痛' in symptom or '腹痛' in symptom for symptom in main_symptoms) and self.patient_profile.get('basic_info', {}).get('性别') == '女':
            return 'B2_腹痛'
        elif any('右上腹痛' in symptom for symptom in main_symptoms):
            return 'B3_右上腹痛'
        elif any('发热' in symptom or '咳嗽' in symptom for symptom in main_symptoms) and isinstance(age, (int, float)) and age <= 12:
            return 'B4_发热咳嗽'
        
        return None
    
    def evaluate_case_specific(self, user_input, case_type):
        """病例特异性评分"""
        score_earned = 0.0
        feedback_items = []
        
        case_criteria = self.case_specific_criteria[case_type]
        
        # 评估关键问诊点
        for point, max_score in case_criteria['关键问诊点'].items():
            if self.check_inquiry_point(user_input, point, case_type):
                sub_score = random.uniform(max_score * 0.8, max_score)
                score_earned += sub_score
                feedback_items.append(f"✓ {point} +{sub_score:.2f}分")
        
        # 评估人文关怀要点
        for point, max_score in case_criteria['人文关怀要点'].items():
            if self.check_humanistic_care(user_input, point, case_type):
                sub_score = random.uniform(max_score * 0.8, max_score)
                score_earned += sub_score
                feedback_items.append(f"✓ {point} +{sub_score:.2f}分")
        
        return score_earned, feedback_items
    
    def check_inquiry_point(self, user_input, point, case_type):
        """检查是否涉及特定问诊点"""
        inquiry_keywords = {
            'A1_腹泻': {
                '工作压力与症状关系': ['工作', '压力', '紧张', '焦虑', '关系'],
                '大便性状描述': ['大便', '性状', '水样', '糊状', '黏液'],
                '发作频率和持续时间': ['频率', '次数', '持续', '时间', '发作'],
                '伴随症状(脐周痛、肠鸣)': ['脐周', '肚脐', '肠鸣', '腹痛'],
                '既往检查结果': ['检查', '化验', '结果', '报告'],
                '睡眠和焦虑状况': ['睡眠', '失眠', '焦虑', '紧张']
            },
            'A2_咯血': {
                '痰血性状和诱因': ['痰', '血', '性状', '诱因', '劳累'],
                '胸痛特点': ['胸痛', '疼痛', '部位', '性质'],
                '吸烟史详细询问': ['吸烟', '抽烟', '烟龄', '戒烟'],
                '家族肺癌史': ['家族', '肺癌', '父亲', '母亲'],
                '既往治疗效果': ['治疗', '效果', '中药', '用药']
            },
            'A3_尿路感染': {
                '尿路症状三联征': ['尿频', '尿急', '尿痛', '排尿'],
                '诱发因素(劳累熬夜)': ['劳累', '熬夜', '装修', '诱因'],
                '月经史和性生活史': ['月经', '性生活', '避孕'],
                '既往类似发作': ['以前', '类似', '发作', '既往'],
                '用药史和效果': ['用药', '氟哌酸', '效果', '治疗']
            },
            'A4_胸痛': {
                '胸痛诱发因素变化': ['诱发', '寒冷', '情绪', '饱餐', '爬楼'],
                '疼痛性质和放射': ['压榨', '放射', '左肩', '性质'],
                '硝酸甘油效果变化': ['硝酸甘油', '含服', '缓解', '效果'],
                '高胆固醇血症史': ['胆固醇', '血脂', '既往'],
                '家族心血管病史': ['家族', '高血压', '心脏病']
            },
            'B1_腰痛': {
                '疼痛性质和诱因': ['腰痛', '性质', '诱因', '提水桶'],
                '体位关系': ['体位', '久坐', '平卧', '姿势'],
                '既往影像学检查': ['MRI', '影像', '检查', '腰椎'],
                '职业相关因素': ['职业', '久坐', '工作'],
                '治疗效果评估': ['治疗', '效果', '护腰', '牵引']
            },
            'B2_腹痛': {
                '疼痛性质和部位': ['腹痛', '撕裂', '部位', '左下腹'],
                '月经史详细询问': ['月经', '末次', '停经', '天数'],
                '阴道流血情况': ['阴道', '流血', '出血', '暗红'],
                '妊娠史和避孕史': ['怀孕', '流产', '避孕', '妊娠'],
                '伴随症状': ['恶心', '呕吐', '肛门', '坠胀']
            },
            'B3_右上腹痛': {
                '疼痛诱因(油腻饮食)': ['油腻', '饮食', '红烧肉', '诱因'],
                '疼痛性质和放射': ['绞痛', '放射', '右肩胛', '性质'],
                '发热和寒战': ['发热', '寒战', '体温'],
                '既往检查结果': ['B超', '胆囊', '结石', '检查'],
                '家族胆石症史': ['家族', '胆结石', '母亲']
            },
            'B4_发热咳嗽': {
                '发热特点和用药': ['发热', '体温', '美林', '退烧药'],
                '咳嗽性质和痰液': ['咳嗽', '痰', '痰鸣', '性质'],
                '喂养和发育史': ['喂养', '发育', '母乳', '辅食'],
                '接触史和诱因': ['接触', '公园', '受凉', '诱因'],
                '伴随症状': ['流涕', '喷嚏', '干呕', '皮疹']
            }
        }
        
        if case_type in inquiry_keywords and point in inquiry_keywords[case_type]:
            keywords = inquiry_keywords[case_type][point]
            return any(keyword in user_input for keyword in keywords)
        
        return False
    
    def check_humanistic_care(self, user_input, point, case_type):
        """检查是否体现人文关怀"""
        care_keywords = {
            'A1_腹泻': {
                '回应癌症担忧': ['癌症', '肿瘤', '担心', '不用担心', '可能性低'],
                '解释肠镜必要性': ['肠镜', '检查', '必要', '排除', '确诊'],
                '缓解检查恐惧': ['不痛', '安全', '放心', '恐惧', '害怕']
            },
            'A2_咯血': {
                '戒烟建议和指导': ['戒烟', '不要抽烟', '吸烟有害', '建议'],
                '肺癌筛查解释': ['筛查', 'CT', '胸片', '检查', '排除'],
                '家族史相关安慰': ['家族史', '不一定', '遗传', '安慰']
            },
            'A3_尿路感染': {
                '解释感染原因': ['感染', '细菌', '原因', '劳累'],
                '生活指导建议': ['多喝水', '休息', '生活', '建议'],
                '检查必要性说明': ['尿检', '检查', '必要', '确诊']
            },
            'A4_胸痛': {
                '心梗风险解释': ['心梗', '风险', '危险', '解释'],
                '紧急检查安排': ['心电图', '心肌酶', '紧急', '检查'],
                '治疗方案讨论': ['治疗', '方案', '支架', '讨论']
            },
            'B1_腰痛': {
                '疾病预后解释': ['预后', '治愈', '恢复', '解释'],
                '生活方式指导': ['生活', '久坐', '锻炼', '指导'],
                '治疗方案制定': ['治疗', '方案', '制定', '计划']
            },
            'B2_腹痛': {
                '宫外孕风险解释': ['宫外孕', '风险', '危险', '解释'],
                '紧急处理安排': ['紧急', '输液', '手术', '处理'],
                '心理支持': ['安慰', '支持', '不要紧张', '放心']
            },
            'B3_右上腹痛': {
                '手术必要性解释': ['手术', '必要', '解释', '治疗'],
                '避免盲目止痛': ['不要止痛', '先检查', '盲目', '止痛'],
                '进一步检查安排': ['进一步', '检查', '安排', 'CT']
            },
            'B4_发热咳嗽': {
                '发热与智力关系解释': ['发热', '智力', '影响', '解释', '不会'],
                '合理用药指导': ['用药', '合理', '指导', '不要乱用'],
                '家长焦虑缓解': ['不要担心', '焦虑', '缓解', '安慰']
            }
        }
        
        if case_type in care_keywords and point in care_keywords[case_type]:
            keywords = care_keywords[case_type][point]
            return any(keyword in user_input for keyword in keywords)
        
        return False
    
    def evaluate_response(self, user_input, context):
        """评估用户回答（更细粒度的评分）"""
        score_earned = 0.0
        feedback_items = []
        
        # 获取当前病例类型
        current_case = self.get_current_case_type()
        
        # 1. 问候礼仪评分
        if context.get("is_greeting"):
            if any(word in user_input for word in ["你好", "您好", "早上好", "下午好", "欢迎"]):
                sub_score = random.uniform(1.8, 2.2)  # 添加浮动使评分更真实
                self.detailed_scores["greeting"]["礼貌称呼"] += sub_score
                feedback_items.append(f"✓ 礼貌称呼 +{sub_score:.2f}分")
                
            if any(word in user_input for word in ["我是", "医生", "主治", "负责"]):
                sub_score = random.uniform(1.8, 2.2)
                self.detailed_scores["greeting"]["自我介绍"] += sub_score
                feedback_items.append(f"✓ 自我介绍 +{sub_score:.2f}分")
                
            if any(word in user_input for word in ["请坐", "放松", "别紧张", "请讲"]):
                sub_score = random.uniform(2.7, 3.3)
                self.detailed_scores["greeting"]["问候患者"] += sub_score
                feedback_items.append(f"✓ 问候患者 +{sub_score:.2f}分")
        
        # 病例特异性评分
        if current_case and current_case in self.case_specific_criteria:
            case_score, case_feedback = self.evaluate_case_specific(user_input, current_case)
            score_earned += case_score
            feedback_items.extend(case_feedback)
        
        # 2. 症状询问评分
        symptom_keywords = {
            "症状持续时间": ["多久", "什么时候", "时间", "持续"],
            "症状性质": ["怎么疼", "怎么样", "感觉", "性质", "描述"],
            "症状程度": ["严重", "程度", "几分", "评级"],
            "症状变化": ["变化", "加重", "减轻", "时好时坏"],
            "伴随症状": ["还有", "伴随", "其他"],
            "诱发缓解因素": ["什么情况", "诱发", "缓解", "什么时候缓解"],
            "与进食关系": ["吃饭", "进食", "饿", "饱"],
            "与体位关系": ["躺", "站", "坐", "姿势", "体位"]
        }
        
        for subcategory, keywords in symptom_keywords.items():
            if any(word in user_input for word in keywords):
                # 根据匹配的关键词数量和质量计算分数
                matched_count = sum(1 for word in keywords if word in user_input)
                max_count = len(keywords)
                match_ratio = min(1.0, matched_count / 3)  # 最多匹配3个就有满分
                sub_score = random.uniform(3.0 * match_ratio * 0.9, 3.0 * match_ratio * 1.1)
                
                # 不能超过子项满分
                max_subscore = self.scoring_criteria["symptom_inquiry"]["子项"][subcategory]
                current_score = self.detailed_scores["symptom_inquiry"][subcategory]
                if current_score + sub_score > max_subscore:
                    sub_score = max_subscore - current_score
                
                if sub_score > 0.01:  # 避免过小的分数
                    self.detailed_scores["symptom_inquiry"][subcategory] += sub_score
                    score_earned += sub_score
                    feedback_items.append(f"✓ {subcategory} +{sub_score:.2f}分")
        
        # 3. 病史询问评分
        history_keywords = {
            "既往疾病史": ["病过", "得病", "患病", "肝炎", "胆结石"],
            "手术史": ["手术", "开刀", "做过手术"],
            "家族史": ["家族", "家人", "兄弟姐妹", "父母"],
            "过敏史": ["过敏", "过敏史", "过敏过"],
            "药物史": ["吃药", "用药", "服药", "吃过的药"],
            "生活史": ["抽烟", "喝酒", "抽烟量", "酒量"]
        }
        
        for subcategory, keywords in history_keywords.items():
            if any(word in user_input for word in keywords):
                matched_count = sum(1 for word in keywords if word in user_input)
                match_ratio = min(1.0, matched_count / 2)  # 最多匹配2个就有满分
                sub_score = random.uniform(2.0 * match_ratio * 0.9, 2.0 * match_ratio * 1.1)
                
                max_subscore = self.scoring_criteria["history_inquiry"]["子项"][subcategory]
                current_score = self.detailed_scores["history_inquiry"][subcategory]
                if current_score + sub_score > max_subscore:
                    sub_score = max_subscore - current_score
                
                if sub_score > 0.01:
                    self.detailed_scores["history_inquiry"][subcategory] += sub_score
                    score_earned += sub_score
                    feedback_items.append(f"✓ {subcategory} +{sub_score:.2f}分")
        
        # 4. 诊断准确性评分
        diagnosis_keywords = {
            "鉴别诊断": ["排除", "鉴别", "可能是别的", "区分"],
            "体征检查": ["摸", "触诊", "叩诊", "听诊"],
            "辅助检查建议": ["检查", "化验", "做CT", "查血"],
            "初步诊断": ["诊断", "可能性最大", "应该是", "考虑"]
        }
        
        for subcategory, keywords in diagnosis_keywords.items():
            if any(word in user_input for word in keywords):
                matched_count = sum(1 for word in keywords if word in user_input)
                match_ratio = min(1.0, matched_count / 2)  # 最多匹配2个就有满分
                sub_score = random.uniform(2.5 * match_ratio * 0.9, 2.5 * match_ratio * 1.1)
                
                max_subscore = self.scoring_criteria["diagnosis_accuracy"]["子项"][subcategory]
                current_score = self.detailed_scores["diagnosis_accuracy"][subcategory]
                if current_score + sub_score > max_subscore:
                    sub_score = max_subscore - current_score
                
                if sub_score > 0.01:
                    self.detailed_scores["diagnosis_accuracy"][subcategory] += sub_score
                    score_earned += sub_score
                    feedback_items.append(f"✓ {subcategory} +{sub_score:.2f}分")
        
        # 5. 治疗计划评分
        treatment_keywords = {
            "治疗解释": ["治疗", "方案", "怎么做", "办法"],
            "手术建议": ["手术", "开刀", "切除", "做手术"],
            "药物建议": ["药物", "药", "化疗", "用药"],
            "随访计划": ["复查", "随访", "定期", "来检查"],
            "生活建议": ["生活", "饮食", "注意", "休息"]
        }
        
        for subcategory, keywords in treatment_keywords.items():
            if any(word in user_input for word in keywords):
                matched_count = sum(1 for word in keywords if word in user_input)
                match_ratio = min(1.0, matched_count / 2)  # 最多匹配2个就有满分
                sub_score = random.uniform(5.0 * match_ratio * 0.9, 5.0 * match_ratio * 1.1)
                
                max_subscore = self.scoring_criteria["treatment_plan"]["子项"][subcategory]
                current_score = self.detailed_scores["treatment_plan"][subcategory]
                if current_score + sub_score > max_subscore:
                    sub_score = max_subscore - current_score
                
                if sub_score > 0.01:
                    self.detailed_scores["treatment_plan"][subcategory] += sub_score
                    score_earned += sub_score
                    feedback_items.append(f"✓ {subcategory} +{sub_score:.2f}分")
        
        # 6. 新增：关键问题检查（扣分项）
        missed_criteria = []
        
        # 检查不当行为和言论的扣分项
        if not hasattr(self, 'deductions'):
            self.deductions = {}
        
        # 检查辱骂和不当言论
        abuse_keywords = ["笨蛋", "白痴", "傻子", "废物", "垃圾", "滚", "闭嘴", "蠢货", "混蛋", "该死", "去死", "妈的", "操", "草", "艹"]
        rude_keywords = ["你懂什么", "别废话", "快点", "烦死了", "别磨蹭", "有完没完", "够了", "别说了"]
        dismissive_keywords = ["没事", "小问题", "不用担心", "想多了", "正常的", "别紧张", "没什么大不了"]
        
        # 辱骂患者扣分
        if any(word in user_input for word in abuse_keywords):
            if "辱骂患者" not in self.deductions:
                deduction = self.scoring_criteria.get("critical_miss", {}).get("扣分项", {}).get("辱骂患者", -15)
                self.current_score += deduction
                self.deductions["辱骂患者"] = deduction
                missed_criteria.append("使用辱骂性语言")
                # 更新患者情绪为愤怒
                self.update_emotion("doctor_abuse")
        
        # 态度恶劣扣分
        elif any(word in user_input for word in rude_keywords):
            if "态度恶劣" not in self.deductions:
                deduction = self.scoring_criteria.get("critical_miss", {}).get("扣分项", {}).get("态度恶劣", -10)
                self.current_score += deduction
                self.deductions["态度恶劣"] = deduction
                missed_criteria.append("医生态度恶劣")
                # 更新患者情绪为愤怒
                self.update_emotion("doctor_rude")
        
        # 不当言论扣分（过于轻视患者症状）
        elif any(word in user_input for word in dismissive_keywords) and len([w for w in dismissive_keywords if w in user_input]) >= 2:
            if "不当言论" not in self.deductions:
                deduction = self.scoring_criteria.get("critical_miss", {}).get("扣分项", {}).get("不当言论", -6)
                self.current_score += deduction
                self.deductions["不当言论"] = deduction
                missed_criteria.append("过于轻视患者症状")
                # 更新患者情绪为沮丧
                self.update_emotion("doctor_dismissive")
        
        # 原有的关键问题检查
        for question, keywords in self.core_questions.items():
            if not any(keyword in user_input for keyword in keywords):
                # 只在相关情境下扣分
                if question == "黄疸" and "黄疸" in str(self.patient_profile.get("symptoms", "")):
                    if "未问黄疸" not in self.deductions:
                        deduction = self.scoring_criteria.get("critical_miss", {}).get("扣分项", {}).get("未问黄疸", -3)
                        self.current_score += deduction
                        self.deductions["未问黄疸"] = deduction
                        missed_criteria.append("未询问黄疸情况")
                
                elif question == "体重变化" and "体重减轻" in str(self.patient_profile.get("symptoms", "")):
                    if "未问体重变化" not in self.deductions:
                        deduction = self.scoring_criteria.get("critical_miss", {}).get("扣分项", {}).get("未问体重变化", -2)
                        self.current_score += deduction
                        self.deductions["未问体重变化"] = deduction
                        missed_criteria.append("未询问体重变化情况")
        
        # 7. 新增：诊断准确性评分（高级评分项，使用独立键避免覆盖基础诊断分数）
        advanced_diag_keywords = {
            "正确识别可切除性": ["可切除", "不可切除", "手术指征", "切除条件"],
            "建议恰当影像检查": ["CT", "MRI", "PET", "超声内镜", "ERCP"],
            "鉴别诊断全面性": ["胆管癌", "慢性胰腺炎", "胰腺囊肿", "十二指肠癌"],
            "肿瘤标志物应用": ["CA19-9", "CEA", "肿瘤标志物", "血清学检查"],
            "分期准确性": ["TNM分期", "局部晚期", "转移", "分期"]
        }
        
        for subcategory, keywords in advanced_diag_keywords.items():
            if any(word in user_input for word in keywords):
                matched_count = sum(1 for word in keywords if word in user_input)
                match_ratio = min(1.0, matched_count / 2)
                sub_score = random.uniform(3.0 * match_ratio * 0.9, 3.0 * match_ratio * 1.1)
                
                # 设置最大子项分数
                max_subscore = 5.0  # 每个子项最高5分
                current_score = self.detailed_scores["advanced_diagnosis"].get(subcategory, 0.0)
                if current_score + sub_score > max_subscore:
                    sub_score = max_subscore - current_score
                
                if sub_score > 0.01:
                    self.detailed_scores["advanced_diagnosis"][subcategory] += sub_score
                    score_earned += sub_score
                    feedback_items.append(f"✓ {subcategory} +{sub_score:.2f}分")
        
        # 添加遗漏关键问题的反馈
        if missed_criteria:
            feedback_items.extend([f"⚠ {criteria}" for criteria in missed_criteria])
        
        self.current_score += score_earned
        self.current_score = min(100.0, max(0.0, self.current_score))  # 确保分数在0-100范围内
        self.score_history.append(self.current_score)
        
        feedback_str = " | ".join(feedback_items) if feedback_items else "无新得分项"
        self.feedback.append(feedback_str)
        
        return score_earned, feedback_str
    
    def get_current_detailed_scores(self):
        """获取当前详细的分数"""
        detailed = {}
        for category in self.scoring_criteria:
            if "子项" in self.scoring_criteria[category]:
                detailed[category] = {
                    "total": sum(self.detailed_scores.get(category, {}).values()),
                    "max": self.scoring_criteria[category]["权重"],
                    "subscores": self.detailed_scores.get(category, {}).copy()
                }
        return detailed
    
    def generate_dynamic_score_chart(self):
        """生成动态评分图表"""
        fig = go.Figure()
        
        if not self.score_history:
            # 空数据时返回空图表
            fig.update_layout(
                title='得分动态变化（暂无数据）',
                xaxis_title='对话轮次',
                yaxis_title='分数',
                yaxis=dict(range=[0, 100])
            )
            return fig
        
        # 添加当前分数折线
        fig.add_trace(go.Scatter(
            x=list(range(1, len(self.score_history)+1)),
            y=self.score_history,
            mode='lines+markers',
            name='当前分数',
            line=dict(color='blue', width=2)
        ))
        
        # 添加目标分数线
        fig.add_trace(go.Scatter(
            x=[0.5, len(self.score_history)+0.5],
            y=[70, 70],
            mode='lines',
            name='合格线',
            line=dict(color='orange', dash='dash')
        ))
        
        # 添加最佳表现线
        fig.add_trace(go.Scatter(
            x=[0.5, len(self.score_history)+0.5],
            y=[90, 90],
            mode='lines',
            name='优秀线',
            line=dict(color='green', dash='dash')
        ))
        
        fig.update_layout(
            title='得分动态变化',
            xaxis_title='对话轮次',
            yaxis_title='分数',
            yaxis=dict(range=[0, 100]),
            showlegend=True
        )
        
        return fig
    
    def generate_score_distribution(self):
        """生成分数分布饼图"""
        detailed = self.get_current_detailed_scores()
        categories = []
        scores = []
        for category, data in detailed.items():
            categories.append(category)
            scores.append(data["total"])
        
        # 创建饼图
        fig = px.pie(
            values=scores,
            names=categories,
            title='能力分布',
            hole=0.3,
            color_discrete_sequence=px.colors.qualitative.Pastel
        )
        
        return fig
    
    def generate_subscore_radar(self):
        """生成各子项分数雷达图"""
        categories = list(self.scoring_criteria.keys())
        subcategories = []
        scores = []
        max_scores = []
        
        # 收集所有子项数据
        for category in categories:
            if "子项" in self.scoring_criteria[category]:
                for subcategory, max_score in self.scoring_criteria[category]["子项"].items():
                    subcategories.append(f"{category} - {subcategory}")
                    scores.append(self.detailed_scores.get(category, {}).get(subcategory, 0))
                    max_scores.append(max_score)
        
        # 创建雷达图
        fig = go.Figure()
        
        # 用户得分
        fig.add_trace(go.Scatterpolar(
            r=scores,
            theta=subcategories,
            fill='toself',
            name='当前得分'
        ))
        
        # 满分线
        fig.add_trace(go.Scatterpolar(
            r=max_scores,
            theta=subcategories,
            name='满分',
            line=dict(color='red', dash='dot')
        ))
        
        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True,
                    range=[0, max(max_scores) * 1.1] if max_scores else [0, 10]
                )),
            showlegend=True,
            title='能力详情分析'
        )
        
        return fig
    
    def generate_report(self):
        """生成练习报告"""
        from docx import Document
        import tempfile
        
        doc = Document()
        doc.add_heading('胰腺外科诊断练习报告', 0)
        
        # 成绩概览
        doc.add_heading('成绩概览', level=1)
        doc.add_paragraph(f"最终得分: {self.current_score:.2f}/100")
        
        # 评分项分析
        doc.add_heading('评分项分析', level=1)
        scoring_data = self.get_current_detailed_scores()
        
        for category, data in scoring_data.items():
            doc.add_paragraph(
                f"{category}: {data['total']:.2f}/{data['max']:.1f}分"
            )
            for subcategory, score in data["subscores"].items():
                max_score = self.scoring_criteria[category]["子项"][subcategory]
                doc.add_paragraph(f"  - {subcategory}: {score:.2f}/{max_score}分")
        
        # 扣分项分析
        doc.add_heading('扣分项分析', level=1)
        for subcategory, deduction in self.deductions.items():
            if deduction != 0:
                doc.add_paragraph(f"  - {subcategory}: {deduction:.1f}分")
        
        # 详细反馈
        doc.add_heading('详细反馈', level=1)
        for i, fb in enumerate(self.feedback, 1):
            doc.add_paragraph(f"第{i}轮: {fb}")
        
        # 学习建议
        doc.add_heading('学习建议', level=1)
        suggestions = [
            "1. 重点复习胰腺癌的典型症状：无痛性黄疸、上腹痛、体重下降",
            "2. 掌握Whipple手术的适应症和手术步骤",
            "3. 熟悉常用化疗方案：FOLFIRINOX、吉西他滨+白蛋白紫杉醇",
            "4. 了解胰腺癌的新辅助治疗和靶向治疗进展",
            "5. 学习术后并发症的诊断和处理：胰瘘、出血等"
        ]
        
        for suggestion in suggestions:
            doc.add_paragraph(suggestion)
        
        # 可视化图表 - 能力分布饼图
        try:
            fig_pie = self.generate_score_distribution()
            doc.add_paragraph("能力分布:")
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_pie:
                fig_pie.write_image(tmp_pie.name, format='png', width=600, height=450)
                doc.add_picture(tmp_pie.name, width=140)
        except Exception as e:
            doc.add_paragraph(f"无法生成能力分布图: {str(e)}")
            doc.add_paragraph("请安装kaleido引擎: pip install kaleido")
        
        # 可视化图表 - 雷达图
        try:
            fig_radar = self.generate_subscore_radar()
            doc.add_paragraph("能力详情分析:")
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_radar:
                fig_radar.write_image(tmp_radar.name, format='png', width=700, height=500)
                doc.add_picture(tmp_radar.name, width=150)
        except Exception as e:
            doc.add_paragraph(f"无法生成能力分析图: {str(e)}")
            doc.add_paragraph("请安装kaleido引擎: pip install kaleido")
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            return tmp.name

# 全局变量
kb = KnowledgeBase()

# AI助手响应生成
def generate_ai_response(question, context, mode="consultation"):
    """使用火山引擎生成AI响应"""
    system_prompt = ""
    
    if mode == "consultation":
        system_prompt = f"""你是一名资深胰腺外科专家，为医生提供专业咨询。请遵循以下准则：
1. 基于最新医学证据（2023-2024 NCCN/ASCO/ESMO指南）回答
2. 区分推荐等级（1级/2A级/2B级/3级）
3. 标注信息来源和量化数据
4. 手术方案注明关键步骤，药物注明剂量和副作用
5. 遇到不确定问题应说明缺乏高质量循证依据
6. 回答简洁专业，不超过300字

当前问题背景：{context}"""
    else:
        system_prompt = f"""你是一名胰腺癌患者，根据以下临床特征扮演患者角色：
{context}

扮演要求：
1. 只描述症状，不提供诊断
2. 逐步提供更多信息，不要一次性透露所有症状
3. 用患者口吻描述，如"医生，我感觉..."
4. 每次回答不超过100字"""
    
    try:
        response = call_ark_api(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question},
            ],
            temperature=0.2,
            max_tokens=500,
            timeout=180,
        )
        return response
    except Exception as e:
        print(f"API调用失败: {str(e)}")
        if mode == "consultation":
            return "我目前无法提供专业咨询，请稍后再试。"
        else:
            return "我觉得不太舒服，但不知道怎么描述。"

# API路由
@app.route('/api/auth/register', methods=['POST'])
def register():
    """用户注册"""
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')
    
    if not username or not password:
        return jsonify({'success': False, 'message': '用户名和密码不能为空'}), 400
    
    try:
        conn = sqlite3.connect('medical_records.db')
        c = conn.cursor()
        
        c.execute("SELECT * FROM yiliao_user WHERE username = ?", (username,))
        if c.fetchone():
            return jsonify({'success': False, 'message': '用户名已存在'}), 400
        
        hashed_pwd = hash_password(password)
        c.execute('''INSERT INTO yiliao_user (username, password) 
                    VALUES (?, ?)''', (username, hashed_pwd))
        conn.commit()
        return jsonify({'success': True, 'message': '注册成功'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'注册失败: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/auth/login', methods=['POST'])
def login():
    """用户登录"""
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')
    
    if not username or not password:
        return jsonify({'success': False, 'message': '用户名和密码不能为空'}), 400
    
    try:
        conn = sqlite3.connect('medical_records.db')
        c = conn.cursor()
        
        hashed_pwd = hash_password(password)
        c.execute("SELECT id FROM yiliao_user WHERE username = ? AND password = ?", 
                 (username, hashed_pwd))
        user = c.fetchone()
        
        if user:
            session['user_id'] = user[0]
            session['username'] = username
            return jsonify({'success': True, 'message': '登录成功', 'username': username})
        else:
            return jsonify({'success': False, 'message': '用户名或密码错误'}), 401
    except Exception as e:
        return jsonify({'success': False, 'message': f'登录失败: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/auth/logout', methods=['POST'])
def logout():
    """用户登出"""
    session.clear()
    return jsonify({'success': True, 'message': '登出成功'})

@app.route('/api/auth/check', methods=['GET'])
def check_auth():
    """检查登录状态"""
    if 'user_id' in session:
        return jsonify({'authenticated': True, 'username': session.get('username')})
    else:
        return jsonify({'authenticated': False})

@app.route('/api/consultation/chat', methods=['POST'])
def consultation_chat():
    """病症咨询聊天"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    data = request.get_json()
    question = data.get('question')
    
    if not question:
        return jsonify({'success': False, 'message': '问题不能为空'}), 400
    
    try:
        # 查询知识库获取上下文
        context = kb.query_knowledge(question)
        
        # 生成AI响应
        response = generate_ai_response(question, context)
        
        # 保存聊天记录到数据库
        conn = sqlite3.connect('medical_records.db')
        c = conn.cursor()
        
        user_id = session['user_id']
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 保存用户问题
        c.execute('''INSERT INTO yiliao_jilu (user_id, timestamp, role, content, mode)
                    VALUES (?, ?, ?, ?, ?)''',
                (user_id, timestamp, 'user', question, 'consultation'))
        
        # 保存AI回答
        c.execute('''INSERT INTO yiliao_jilu (user_id, timestamp, role, content, mode)
                    VALUES (?, ?, ?, ?, ?)''',
                (user_id, timestamp, 'assistant', response, 'consultation'))
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True, 
            'response': response,
            'context': context
        })
    except Exception as e:
        return jsonify({'success': False, 'message': f'处理失败: {str(e)}'}), 500

@app.route('/api/consultation/history', methods=['GET'])
def get_consultation_history():
    """获取咨询历史"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    try:
        conn = sqlite3.connect('medical_records.db')
        c = conn.cursor()
        
        user_id = session['user_id']
        c.execute('''SELECT role, content, timestamp FROM yiliao_jilu 
                    WHERE user_id = ? AND mode = 'consultation' 
                    ORDER BY timestamp ASC''', (user_id,))
        
        records = c.fetchall()
        history = []
        for record in records:
            history.append({
                'role': record[0],
                'content': record[1],
                'timestamp': record[2]
            })
        
        conn.close()
        return jsonify({'success': True, 'history': history})
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取历史失败: {str(e)}'}), 500

@app.route('/api/consultation/clear', methods=['POST'])
def clear_consultation_history():
    """清除咨询历史"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    try:
        conn = sqlite3.connect('medical_records.db')
        c = conn.cursor()
        
        user_id = session['user_id']
        c.execute('''DELETE FROM yiliao_jilu 
                    WHERE user_id = ? AND mode = 'consultation' ''', (user_id,))
        
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'message': '历史记录已清除'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'清除失败: {str(e)}'}), 500

@app.route('/api/practice/start', methods=['POST'])
def start_practice():
    """开始练习"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    try:
        # 创建新的练习系统实例
        practice_system = PracticeSystem(kb)
        
        # 将练习系统存储在session中（实际应用中可能需要使用Redis等）
        session['practice_system'] = {
            'patient_profile': practice_system.patient_profile,
            'patient_personality': practice_system.patient_personality,
            'emotion_state': practice_system.emotion_state,
            'current_score': practice_system.current_score,
            'score_history': practice_system.score_history,
            'detailed_scores': practice_system.detailed_scores,
            'feedback': practice_system.feedback
        }
        
        return jsonify({
            'success': True,
            'patient_profile': practice_system.patient_profile,
            'patient_personality': practice_system.patient_personality,
            'emotion_state': practice_system.emotion_state,
            'current_score': practice_system.current_score
        })
    except Exception as e:
        return jsonify({'success': False, 'message': f'开始练习失败: {str(e)}'}), 500

@app.route('/api/practice/chat', methods=['POST'])
def practice_chat():
    """练习模式聊天"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    if 'practice_system' not in session:
        return jsonify({'success': False, 'message': '请先开始练习'}), 400
    
    data = request.get_json()
    user_input = data.get('input')
    is_greeting = data.get('is_greeting', False)
    
    if not user_input:
        return jsonify({'success': False, 'message': '输入不能为空'}), 400
    
    try:
        # 重建练习系统实例
        practice_system = PracticeSystem(kb)
        practice_data = session['practice_system']
        
        # 恢复状态
        practice_system.patient_profile = practice_data['patient_profile']
        practice_system.patient_personality = practice_data['patient_personality']
        practice_system.emotion_state = practice_data['emotion_state']
        practice_system.current_score = practice_data['current_score']
        practice_system.score_history = practice_data['score_history']
        practice_system.detailed_scores = practice_data['detailed_scores']
        practice_system.feedback = practice_data['feedback']
        practice_system.deductions = practice_data.get('deductions', {})
        
        # 评估用户回答
        print(f"[SCORE-DEBUG] 评估前总分: {practice_system.current_score}, deductions: {practice_system.deductions}")
        score_earned, feedback = practice_system.evaluate_response(
            user_input, 
            {
                "is_greeting": is_greeting,
                "symptoms": practice_system.patient_profile["symptoms"]
            }
        )
        print(f"[SCORE-DEBUG] 本轮得分: {score_earned}, 评估后总分: {practice_system.current_score}, feedback: {feedback}")
        
        # 生成患者回答
        patient_response = practice_system.generate_patient_response(user_input)
        
        # 保存聊天记录到数据库
        conn = sqlite3.connect('medical_records.db')
        c = conn.cursor()
        
        user_id = session['user_id']
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 保存医生问题
        c.execute('''INSERT INTO yiliao_jilu (user_id, timestamp, role, content, mode, score)
                    VALUES (?, ?, ?, ?, ?, ?)''',
                (user_id, timestamp, 'user', user_input, 'practice', practice_system.current_score))
        
        # 保存患者回答
        c.execute('''INSERT INTO yiliao_jilu (user_id, timestamp, role, content, mode, score)
                    VALUES (?, ?, ?, ?, ?, ?)''',
                (user_id, timestamp, 'assistant', patient_response, 'practice', practice_system.current_score))
        
        conn.commit()
        conn.close()
        
        # 更新session中的练习系统状态
        session['practice_system'] = {
            'patient_profile': practice_system.patient_profile,
            'patient_personality': practice_system.patient_personality,
            'emotion_state': practice_system.emotion_state,
            'current_score': practice_system.current_score,
            'score_history': practice_system.score_history,
            'detailed_scores': practice_system.detailed_scores,
            'feedback': practice_system.feedback,
            'deductions': practice_system.deductions
        }
        
        return jsonify({
            'success': True,
            'patient_response': patient_response,
            'score_earned': score_earned,
            'feedback': feedback,
            'current_score': practice_system.current_score,
            'emotion_state': practice_system.emotion_state
        })
    except Exception as e:
        return jsonify({'success': False, 'message': f'处理失败: {str(e)}'}), 500

@app.route('/api/practice/charts', methods=['GET'])
def get_practice_charts():
    """获取练习图表数据"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    if 'practice_system' not in session:
        return jsonify({'success': False, 'message': '请先开始练习'}), 400
    
    try:
        # 重建练习系统实例
        practice_system = PracticeSystem(kb)
        practice_data = session['practice_system']
        
        # 恢复状态
        practice_system.current_score = practice_data['current_score']
        practice_system.score_history = practice_data['score_history']
        practice_system.detailed_scores = practice_data['detailed_scores']
        
        # 生成图表
        score_chart_fig = practice_system.generate_dynamic_score_chart()
        distribution_chart_fig = practice_system.generate_score_distribution()
        radar_chart_fig = practice_system.generate_subscore_radar()
        detailed_scores = practice_system.get_current_detailed_scores()
        
        # 转换为JSON格式
        score_chart = pio.to_json(score_chart_fig)
        distribution_chart = pio.to_json(distribution_chart_fig)
        radar_chart = pio.to_json(radar_chart_fig)
        
        return jsonify({
            'success': True,
            'score_chart': score_chart,
            'distribution_chart': distribution_chart,
            'radar_chart': radar_chart,
            'detailed_scores': detailed_scores
        })
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取图表失败: {str(e)}'}), 500

@app.route('/api/practice/report', methods=['GET'])
def generate_practice_report():
    """生成练习报告"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    if 'practice_system' not in session:
        return jsonify({'success': False, 'message': '请先开始练习'}), 400
    
    try:
        # 重建练习系统实例
        practice_system = PracticeSystem(kb)
        practice_data = session['practice_system']
        
        # 恢复状态
        practice_system.current_score = practice_data['current_score']
        practice_system.score_history = practice_data['score_history']
        practice_system.detailed_scores = practice_data['detailed_scores']
        practice_system.deductions = practice_data.get('deductions', {})
        practice_system.feedback = practice_data.get('feedback', [])
        practice_system.scoring_criteria = practice_data.get('scoring_criteria', practice_system.scoring_criteria)
        
        # 生成报告
        report_path = practice_system.generate_report()
        
        # 读取文件并返回
        with open(report_path, 'rb') as f:
            report_data = f.read()
        
        import os
        os.unlink(report_path)  # 删除临时文件
        
        from flask import make_response
        import time
        
        response = make_response(report_data)
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename=胰腺诊断报告_{time.strftime("%Y%m%d")}.docx'
        
        return response
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'生成报告失败: {str(e)}'}), 500

@app.route('/api/practice/end', methods=['POST'])
def end_practice():
    """结束练习"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    if 'practice_system' not in session:
        return jsonify({'success': False, 'message': '没有进行中的练习'}), 400
    
    try:
        practice_data = session['practice_system']
        final_score = practice_data['current_score']
        detailed_scores = practice_data['detailed_scores']
        
        # 重建练习系统实例以生成图表
        practice_system = PracticeSystem(kb)
        practice_system.current_score = practice_data['current_score']
        practice_system.score_history = practice_data['score_history']
        practice_system.detailed_scores = practice_data['detailed_scores']
        
        # 生成图表数据
        score_chart = pio.to_json(practice_system.generate_dynamic_score_chart())
        distribution_chart = pio.to_json(practice_system.generate_score_distribution())
        radar_chart = pio.to_json(practice_system.generate_subscore_radar())
        
        # 清除练习状态
        session.pop('practice_system', None)
        
        return jsonify({
            'success': True,
            'final_score': final_score,
            'detailed_scores': detailed_scores,
            'score_chart': score_chart,
            'distribution_chart': distribution_chart,
            'radar_chart': radar_chart,
            'message': '练习已结束'
        })
    except Exception as e:
        return jsonify({'success': False, 'message': f'结束练习失败: {str(e)}'}), 500

if __name__ == '__main__':
    init_database()
    app.run(debug=False, host='0.0.0.0', port=8877)