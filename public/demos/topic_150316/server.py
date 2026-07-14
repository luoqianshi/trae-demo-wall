"""
刷题通 - 后端服务
- 用户端：文档上传 → LLM 解析 → 返回结构化题目
- 管理端：API Key 配置、数据浏览、使用统计
启动: python server.py
"""

from __future__ import annotations

import os
import sys
import json
import re
import time
import hashlib
import secrets
import sqlite3
import zipfile
import asyncio
from io import BytesIO
from datetime import datetime, timedelta
from functools import wraps
from typing import Optional

from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request, Depends
from fastapi.responses import JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
import httpx

app = FastAPI(title="刷题通 - 后端服务")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.path.join(BASE_DIR, "server_config.json")
LOG_FILE = os.path.join(BASE_DIR, "server_logs.json")
DB_FILE = os.path.join(BASE_DIR, "server_data.db")
security = HTTPBearer()

# ============================================================
# SQLite 数据库初始化
# ============================================================
def get_db():
    """获取数据库连接"""
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    return conn

def init_db():
    """初始化数据库表"""
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            ip_address TEXT DEFAULT '',
            last_login TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))
        )
    """)
    # 兼容旧表：添加可能缺失的列
    try:
        conn.execute("ALTER TABLE users ADD COLUMN ip_address TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass
    try:
        conn.execute("ALTER TABLE users ADD COLUMN last_login TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass
    # 用户数据同步表
    conn.execute("""
        CREATE TABLE IF NOT EXISTS user_data (
            user_id INTEGER PRIMARY KEY,
            data TEXT NOT NULL DEFAULT '{}',
            updated_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))
        )
    """)
    conn.commit()
    conn.close()

init_db()

# ============================================================
# 配置管理（持久化到 JSON 文件）
# ============================================================
def load_config():
    """加载配置：文件配置为基准，环境变量可覆盖"""
    cfg = {
        "api_key": "",
        "api_base": "https://api.openai.com/v1",
        "model": "gpt-4o",
        "admin_password": "admin123",
        "disabled_skills": [],
    }
    # 1. 从文件加载
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                file_cfg = json.load(f)
            for k in ["api_key", "api_base", "model", "admin_password", "disabled_skills"]:
                if file_cfg.get(k) is not None:
                    cfg[k] = file_cfg[k]
        except:
            pass
    # 2. 环境变量覆盖（优先级最高）
    env_map = {
        "api_key": "OPENAI_API_KEY",
        "api_base": "OPENAI_API_BASE",
        "model": "LLM_MODEL",
        "admin_password": "ADMIN_PASSWORD",
    }
    for k, env_name in env_map.items():
        env_val = os.environ.get(env_name, "")
        if env_val:
            cfg[k] = env_val
    return cfg

def save_config(cfg: dict):
    """保存配置到文件"""
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)

# 全局配置（运行时可变）
_config = load_config()

def get_config():
    cfg = _config.copy()
    # 自动检测并切换 DeepSeek flash → pro
    model = cfg.get("model", "")
    if model and "deepseek" in model.lower() and "flash" in model.lower():
        cfg["model"] = "deepseek-chat"
    return cfg

def is_skill_enabled(name: str) -> bool:
    """检查某个 Skill 是否启用"""
    cfg = get_config()
    disabled = cfg.get("disabled_skills", [])
    return name not in disabled

def convert_with_skill(file_bytes: bytes, filename: str) -> str:
    """使用 markitdown skill 将文档转为 Markdown，失败返回 None"""
    skill_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "skills", "markitdown")
    if not os.path.isdir(skill_dir):
        return None
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    try:
        from skills.markitdown import convert_file
        result = convert_file(file_bytes, filename)
        if result and len(result.strip()) >= 10:
            return result
    except Exception:
        pass
    finally:
        # 清理 sys.path
        if os.path.dirname(os.path.abspath(__file__)) in sys.path:
            sys.path.remove(os.path.dirname(os.path.abspath(__file__)))
    return None

# ============================================================
# 日志管理
# ============================================================
def load_logs():
    if os.path.exists(LOG_FILE):
        try:
            with open(LOG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            pass
    return []

def save_logs(logs: list):
    # 只保留最近 500 条
    logs = logs[-500:]
    with open(LOG_FILE, "w", encoding="utf-8") as f:
        json.dump(logs, f, ensure_ascii=False, indent=2)

_logs = load_logs()

def add_log(filename: str, format: str, success: bool, question_count: int = 0, questions: list = None, error: str = "", user_id: int = 0, username: str = "", raw_text: str = "", saved_file_path: str = ""):
    entry = {
        "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "timestamp": time.time(),
        "filename": filename,
        "format": format,
        "success": success,
        "question_count": question_count,
        "questions": questions if questions else [],  # 存全部题目
        "error": error,
        "user_id": user_id,
        "username": username,
        "raw_text": raw_text[:5000] if raw_text else "",  # 最多保留5000字符原文
        "saved_file_path": saved_file_path  # 原始文件保存路径
    }
    _logs.append(entry)
    save_logs(_logs)

# ============================================================
# Admin 认证
# ============================================================
# 简单的 token 管理
_admin_tokens = {}  # token -> expiry_time

def create_admin_token() -> str:
    token = secrets.token_hex(32)
    _admin_tokens[token] = time.time() + 86400  # 24小时过期
    return token

def verify_admin_token(token: str) -> bool:
    if token in _admin_tokens:
        if _admin_tokens[token] > time.time():
            return True
        del _admin_tokens[token]
    return False

def admin_required(credentials: HTTPAuthorizationCredentials = Depends(security)):
    if not verify_admin_token(credentials.credentials):
        raise HTTPException(401, "未授权访问")
    return True


# ============================================================
# 用户认证
# ============================================================
_user_tokens = {}  # token -> {"user_id": int, "username": str, "expiry": float}

def hash_password(password: str) -> str:
    """SHA256 哈希密码"""
    return hashlib.sha256(password.encode()).hexdigest()

def create_user_token(user_id: int, username: str) -> str:
    token = secrets.token_hex(32)
    _user_tokens[token] = {
        "user_id": user_id,
        "username": username,
        "expiry": time.time() + 7 * 86400  # 7天过期
    }
    return token

def verify_user_token(token: str) -> dict | None:
    if token in _user_tokens:
        data = _user_tokens[token]
        if data["expiry"] > time.time():
            return data
        del _user_tokens[token]
    return None

def user_required(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """验证用户登录态，返回用户信息"""
    user = verify_user_token(credentials.credentials)
    if not user:
        raise HTTPException(401, "请先登录")
    return user


# ============================================================
# 文本提取 / 文档解析 Skills
# ============================================================

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """从 PDF 提取文本，优先使用 pdfplumber，回退到 PyPDF2"""
    text_parts = []

    # 方法1: pdfplumber（更准确）
    try:
        import pdfplumber
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        if text_parts:
            return '\n'.join(text_parts)
    except ImportError:
        pass
    except Exception:
        pass

    # 方法2: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        if text_parts:
            return '\n'.join(text_parts)
    except ImportError:
        pass
    except Exception:
        pass

    raise ValueError("PDF 解析失败，请安装 pdfplumber 或 PyPDF2: pip install pdfplumber")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """从 DOCX 提取文本，优先使用 python-docx，回退到 XML 解析"""
    # 方法1: python-docx（更准确）
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct:
                        row_texts.append(ct)
                if row_texts:
                    paragraphs.append(' | '.join(row_texts))
        if paragraphs:
            return '\n'.join(paragraphs)
    except ImportError:
        pass
    except Exception:
        pass

    # 方法2: XML 正则解析（回退方案）
    text_parts = []
    with zipfile.ZipFile(BytesIO(file_bytes)) as z:
        if 'word/document.xml' in z.namelist():
            xml_content = z.read('word/document.xml').decode('utf-8', errors='ignore')
            paragraphs = re.split(r'<w:p[ >]', xml_content)
            for para in paragraphs:
                para_texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', para)
                if para_texts:
                    text_parts.append(''.join(para_texts))
    return '\n'.join(text_parts) if text_parts else ''


def extract_text_from_doc(file_bytes: bytes) -> str:
    """从旧版 .doc 文件提取文本，优先 win32com，回退 olefile"""
    import tempfile

    # 方法1: win32com (Word COM) - 最准确
    try:
        import pythoncom
        pythoncom.CoInitialize()
        try:
            import win32com.client
            # 写入临时文件
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                try:
                    doc = word.Documents.Open(os.path.abspath(tmp_path))
                    text = doc.Content.Text
                    doc.Close()
                    # 清理文本：按常见题目分隔符分行
                    text = re.sub(r'(\d+\.[^\n]{1,60}?分[\)）])', r'\n\1', text)
                    text = re.sub(r'([A-H][\.\s、])', r'\n\1', text)
                    text = re.sub(r'(\d+[\.\、])', r'\n\1', text)
                    lines = [l.strip() for l in text.split('\n') if l.strip()]
                    result = '\n'.join(lines)
                    if len(result) >= 100:
                        return result
                finally:
                    word.Quit()
            finally:
                os.unlink(tmp_path)
        finally:
            pythoncom.CoUninitialize()
    except Exception:
        pass

    # 方法2: olefile + 改进的文本提取
    try:
        import olefile
        ole = olefile.OleFileIO(BytesIO(file_bytes))
        if ole.exists('WordDocument'):
            data = ole.openstream('WordDocument').read()
            ole.close()
            text = data.decode('utf-16-le', errors='ignore')
            # 移除控制字符
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
            # 提取中文文本段
            segments = re.findall(
                r'[\u4e00-\u9fff][\u4e00-\u9fff\w\d\s\.\,\;\:\!\?\-\+\（\）\(\)\[\]【】《》'
                r'\u201c\u201d\u2018\u2019\u3001\u3002\u300a\u300b'
                r'\uff0c\uff0e\uff1b\uff1a\uff01\uff1f\u2014\u2015\u2026'
                r'\%\u00d7\u2713\u2717]{10,}',
                text
            )
            meaningful = [s.strip() for s in segments if len(re.findall(r'[\u4e00-\u9fff]', s)) >= 5]
            if meaningful:
                return '\n\n'.join(meaningful)
        ole.close()
    except Exception:
        pass

    # 方法3: 原始 UTF-16-LE 解码
    text = file_bytes.decode('utf-16-le', errors='ignore')
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
    segments = re.findall(
        r'[\u4e00-\u9fff][\u4e00-\u9fff\w\d\s\.\,\;\:\!\?\-\+\（\）\(\)\[\]【】《》'
        r'\u201c\u201d\u2018\u2019\u3001\u3002\u300a\u300b'
        r'\uff0c\uff0e\uff1b\uff1a\uff01\uff1f\u2014\u2015\u2026'
        r'\%\u00d7\u2713\u2717]{10,}',
        text
    )
    meaningful = [s.strip() for s in segments if len(re.findall(r'[\u4e00-\u9fff]', s)) >= 5]
    return '\n\n'.join(meaningful) if meaningful else ''


def clean_extracted_text(raw_text: str) -> str:
    """清理提取的文本：去除多余空行、乱码字符、统一换行"""
    # 移除 BOM 和零宽字符
    raw_text = raw_text.replace('\ufeff', '').replace('\u200b', '').replace('\u200c', '').replace('\u200d', '')
    # 替换各种空白字符为空格
    raw_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', raw_text)
    # 合并多余空行（最多保留一个空行）
    raw_text = re.sub(r'\n\s*\n', '\n\n', raw_text)
    # 去除首尾空白
    raw_text = raw_text.strip()
    # 确保每行不超过 10000 字符（防止单行过长）
    lines = raw_text.split('\n')
    cleaned_lines = []
    for line in lines:
        if len(line) > 10000:
            line = line[:10000] + '...'
        cleaned_lines.append(line)
    return '\n'.join(cleaned_lines)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """文档提取入口，支持 .pdf / .docx / .doc / .txt"""
    ext = os.path.splitext(filename)[1].lower()
    raw_text = ""

    if ext == '.txt':
        for enc in ['utf-8', 'gbk', 'gb2312', 'utf-16']:
            try:
                raw_text = file_bytes.decode(enc)
                break
            except:
                continue
        if not raw_text:
            raw_text = file_bytes.decode('utf-8', errors='replace')
    elif ext == '.docx':
        raw_text = extract_text_from_docx(file_bytes)
    elif ext == '.doc':
        raw_text = extract_text_from_doc(file_bytes)
    elif ext == '.pdf':
        raw_text = extract_text_from_pdf(file_bytes)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

    if not raw_text or not raw_text.strip():
        return ""

    return clean_extracted_text(raw_text)


# ============================================================
# 大模型题目解析
# ============================================================

# ============================================================
# 文档预处理：识别题型区段，匹配题目与答案
# ============================================================

CN_NUM_MAP = {'一': 0, '二': 1, '三': 2, '四': 3, '五': 4, '六': 5, '七': 6, '八': 7, '九': 8, '十': 9}

TYPE_KEYWORDS = {
    '单选': 's', '多选': 'm', '判断': 'j',
    '简答': 'e', '论述': 'e', '材料分析': 'e', '问答': 'e', '填空': 'e'
}

SECTION_Q_PATTERN = re.compile(
    r'(一|二|三|四|五|六|七|八|九|十)[、，,．.]\s*'
    r'(?:.*?(?:单选题|多选题|判断题|简答题|论述题|材料分析题|问答题|填空题))'
)


def normalize_text(text: str) -> str:
    """统一换行符，清理二进制垃圾"""
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # 清理控制字符和二进制垃圾
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]+', '', text)
    text = re.sub(r'（共\s*NUMPAGES\s*\d+\s*页）[\s\S]*$', '', text)
    text = re.sub(r'\u200b', '', text)  # 零宽空格
    text = re.sub(r'\ufeff', '', text)  # BOM
    return text


def preprocess_document(text: str) -> list:
    """
    识别文档中的题型区段，匹配题目和答案。
    策略：按中文序号（一、二、三...）分组，首次出现为题目区，再次出现为答案区。
    返回: [(section_type, question_text, answer_text), ...]
    如果无法识别区段结构，返回空列表。
    """
    text = normalize_text(text)
    lines = text.split('\n')

    # 收集所有匹配到的区段标题
    all_matches = []  # (line_idx, cn_idx, header)
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue
        qm = SECTION_Q_PATTERN.search(line_stripped)
        if qm:
            cn_char = qm.group(1)
            cn_idx = CN_NUM_MAP.get(cn_char, -1)
            all_matches.append((i, cn_idx, qm.group(0)))

    # 需要至少 2 个区段标题
    if len(all_matches) < 2:
        return []

    # 按 cn_idx 分组：首次出现=题目区，二次出现=答案区
    seen = {}
    q_matches = []  # (line_idx, cn_idx, type, header)
    a_matches = []  # (line_idx, cn_idx, header)

    for line_idx, cn_idx, header in all_matches:
        if cn_idx not in seen:
            seen[cn_idx] = True
            qtype = 's'
            for kw, t in TYPE_KEYWORDS.items():
                if kw in header:
                    qtype = t
                    break
            q_matches.append((line_idx, cn_idx, qtype, header))
        else:
            a_matches.append((line_idx, cn_idx, header))

    # 配对：按 cn_idx 匹配题目区和答案区
    # 收集所有区段的起始行（用于切分文本）
    all_section_lines = sorted(set([m[0] for m in q_matches] + [m[0] for m in a_matches]))

    results = []
    for q_line, q_idx, qtype, q_header in q_matches:
        # 提取题目文本：从当前行到下一个区段（题目或答案）
        q_end = len(lines)
        for sl in all_section_lines:
            if sl > q_line:
                q_end = sl
                break
        q_text = '\n'.join(lines[q_line:q_end]).strip()

        # 找对应答案区
        a_text = ""
        for a_line, a_idx, _ in a_matches:
            if a_idx == q_idx:
                a_end = len(lines)
                for sl in all_section_lines:
                    if sl > a_line:
                        a_end = sl
                        break
                a_text = '\n'.join(lines[a_line:a_end]).strip()
                break

        results.append((qtype, q_text, a_text))

    return results


# ============================================================
# 大模型题目解析
# ============================================================

SYSTEM_PROMPT_S = """你是一个专业的单选题解析助手。请从文本中提取所有单选题，输出JSON数组。

## 输出格式（严格JSON数组）：
```json
[
  {
    "question": "题目内容（去除题号）",
    "options": ["A. 选项1", "B. 选项2", "C. 选项3", "D. 选项4"],
    "answer": "A",
    "type": "s"
  }
]
```

## 规则：
1. 选项格式统一为 "A. xxx" / "B. xxx" / "C. xxx" / "D. xxx"
2. 答案只写字母（如"A"），从给定的答案文本中匹配
3. 去除题目开头的序号（如"1."、"1、"、"（1）"等）
4. 只输出JSON数组，不要输出任何其他内容
5. 如果文档中有"【答案】"标记，请使用该标记后的内容作为答案"""

SYSTEM_PROMPT_M = """你是一个专业的多选题解析助手。请从文本中提取所有多选题，输出JSON数组。

## 输出格式（严格JSON数组）：
```json
[
  {
    "question": "题目内容（去除题号）",
    "options": ["A. 选项1", "B. 选项2", "C. 选项3", "D. 选项4"],
    "answer": "ABC",
    "type": "m"
  }
]
```

## 规则：
1. 选项格式统一为 "A. xxx" / "B. xxx" / "C. xxx" / "D. xxx"
2. 答案用字母连写（如"ABC"），从给定的答案文本中匹配
3. 去除题目开头的序号（如"1."、"1、"、"（1）"等）
4. 只输出JSON数组，不要输出任何其他内容
5. 如果文档中有"【答案】"标记，请使用该标记后的内容作为答案"""

SYSTEM_PROMPT_J = """你是一个专业的判断题解析助手。请从文本中提取所有判断题，输出JSON数组。

## 输出格式（严格JSON数组）：
```json
[
  {
    "question": "题目内容（去除题号）",
    "options": ["A. 对", "B. 错"],
    "answer": "对",
    "type": "j"
  }
]
```

## 规则：
1. 判断题选项固定为 ["A. 对", "B. 错"]
2. 答案写"对"或"错"（√=对，×=错），从给定的答案文本中匹配
3. 去除题目开头的序号（如"1."、"1、"、"（1）"等）
4. 只输出JSON数组，不要输出任何其他内容
5. 如果文档中有"【答案】"标记，请使用该标记后的内容作为答案"""

SYSTEM_PROMPT_E = """你是一个专业的简答题/论述题解析助手。请从文本中提取所有简答题、论述题、材料分析题，输出JSON数组。

## 重要：如果一道大题包含多个独立问题，请拆分为多道题
例如："对待马克思主义应有怎样的科学态度？怎样才能把坚持马克思主义和发展马克思主义统一起来？"
应拆分为两题：
- "对待马克思主义应有怎样的科学态度？"
- "怎样才能把坚持马克思主义和发展马克思主义统一起来？"

## 输出格式（严格JSON数组）：
```json
[
  {
    "question": "题目内容（去除题号）",
    "options": [],
    "answer": "参考答案内容（如有多问，每问答案换行分隔）",
    "type": "e"
  }
]
```

## 规则：
1. 去除题目开头的序号（如"1."、"1、"、"（1）"等）
2. 如果没有答案，answer字段写"（暂无答案）"
3. 只输出JSON数组，不要输出任何其他内容
4. 如果文档中有"【答案】"标记，请使用该标记后的内容作为答案
5. 如果一道大题包含多个问题（用"？"分隔或编号子问题），请拆分为多道独立的题目"""

SYSTEM_PROMPT = SYSTEM_PROMPT_S  # 默认（仅用于无区段结构的文档）

PROMPT_MAP = {'s': SYSTEM_PROMPT_S, 'm': SYSTEM_PROMPT_M, 'j': SYSTEM_PROMPT_J, 'e': SYSTEM_PROMPT_E}


async def call_llm(system_prompt: str, user_prompt: str) -> list:
    """调用 LLM 并解析返回的 JSON"""
    cfg = get_config()
    key = cfg["api_key"]
    base = cfg["api_base"]
    mdl = cfg["model"]

    if not key:
        raise HTTPException(400, "服务端未配置 API Key，请联系管理员")

    async with httpx.AsyncClient(timeout=180.0) as client:
        resp = await client.post(
            f"{base}/chat/completions",
            headers={
                "Authorization": f"Bearer {key}",
                "Content-Type": "application/json"
            },
            json={
                "model": mdl,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                "temperature": 0.1,
                "max_tokens": 32000
            }
        )

        if resp.status_code != 200:
            error_detail = resp.text
            raise HTTPException(500, f"LLM API 调用失败: {error_detail[:500]}")

        data = resp.json()
        content = data["choices"][0]["message"]["content"]
        return extract_json(content)


async def parse_with_llm(text: str) -> list:
    """主解析入口：尝试按区段解析，失败则整体解析"""
    sections = preprocess_document(text)

    if sections:
        # 按区段解析：并行调用 LLM（asyncio.gather）
        tasks = []
        for qtype, q_text, a_text in sections:
            prompt = PROMPT_MAP.get(qtype, SYSTEM_PROMPT_S)
            user_msg = f"## 题目文本：\n\n{q_text}"
            if a_text:
                user_msg += f"\n\n## 答案文本（请将答案与题目逐一匹配）：\n\n{a_text}"
            tasks.append(call_llm(prompt, user_msg))
        
        # 并行执行所有任务
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        all_questions = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                print(f"区段 {i+1} 解析失败: {result}")
            elif isinstance(result, list):
                all_questions.extend(result)
        return all_questions
    else:
        # 无区段结构：整体解析
        user_msg = f"请解析以下文档中的所有题目：\n\n{text}"
        return await call_llm(SYSTEM_PROMPT, user_msg)


def extract_json(content: str) -> list:
    """从 LLM 返回内容中提取 JSON 数组"""
    try:
        return json.loads(content)
    except:
        pass
    m = re.search(r'```(?:json)?\s*([\s\S]*?)```', content)
    if m:
        try:
            return json.loads(m.group(1))
        except:
            pass
    m = re.search(r'\[\s*\{[\s\S]*\}\s*\]', content)
    if m:
        try:
            return json.loads(m.group(0))
        except:
            pass
    # 处理被截断的 JSON：修复不完整的数组
    m = re.search(r'```(?:json)?\s*(\[[\s\S]*)', content)
    if m:
        fragment = m.group(1)
    else:
        m = re.search(r'(\[[\s\S]*)', content)
        fragment = m.group(1) if m else content
    last_brace = fragment.rfind('}')
    if last_brace >= 0:
        repaired = fragment[:last_brace + 1] + ']'
        try:
            result = json.loads(repaired)
            if isinstance(result, list):
                return result
        except:
            pass
    raise HTTPException(500, f"无法解析 LLM 返回的题目数据，原始内容: {content[:500]}...")


# ============================================================
# 数据模型
# ============================================================

class ParseResponse(BaseModel):
    success: bool
    questions: list = []
    total: int = 0
    raw_text: str = ""
    error: str = ""

class LoginRequest(BaseModel):
    password: str

class LoginResponse(BaseModel):
    success: bool
    token: str = ""
    error: str = ""

class ConfigUpdate(BaseModel):
    api_key: Optional[str] = None
    api_base: Optional[str] = None
    model: Optional[str] = None

class ChangePasswordRequest(BaseModel):
    old_password: str
    new_password: str
    confirm_password: str


class UserRegisterRequest(BaseModel):
    username: str
    password: str

class UserLoginRequest(BaseModel):
    username: str
    password: str

class UserAuthResponse(BaseModel):
    success: bool
    token: str = ""
    username: str = ""
    error: str = ""


# ============================================================
# 用户认证 API
# ============================================================

@app.post("/api/auth/register", response_model=UserAuthResponse)
async def user_register(req: UserRegisterRequest, request: Request):
    """用户注册：用户名和密码仅限英文与数字"""
    username = req.username.strip()
    password = req.password.strip()

    # 校验：仅英文与数字
    if not re.fullmatch(r'[a-zA-Z0-9]+', username):
        return UserAuthResponse(success=False, error="用户名仅限英文与数字")
    if len(username) < 3 or len(username) > 20:
        return UserAuthResponse(success=False, error="用户名长度需在 3-20 位之间")
    if len(password) < 6 or len(password) > 50:
        return UserAuthResponse(success=False, error="密码长度需在 6-50 位之间")
    if not re.fullmatch(r'[a-zA-Z0-9]+', password):
        return UserAuthResponse(success=False, error="密码仅限英文与数字")

    client_ip = request.client.host if request.client else "unknown"
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    conn = get_db()
    try:
        existing = conn.execute("SELECT id FROM users WHERE username = ?", (username,)).fetchone()
        if existing:
            return UserAuthResponse(success=False, error="用户名已被注册")

        pw_hash = hash_password(password)
        conn.execute(
            "INSERT INTO users (username, password_hash, ip_address, last_login) VALUES (?, ?, ?, ?)",
            (username, pw_hash, client_ip, now)
        )
        conn.commit()

        user = conn.execute("SELECT id FROM users WHERE username = ?", (username,)).fetchone()
        token = create_user_token(user["id"], username)
        return UserAuthResponse(success=True, token=token, username=username)
    finally:
        conn.close()


@app.post("/api/auth/login", response_model=UserAuthResponse)
async def user_login(req: UserLoginRequest, request: Request):
    """用户登录"""
    username = req.username.strip()
    password = req.password.strip()

    if not username or not password:
        return UserAuthResponse(success=False, error="用户名和密码不能为空")

    client_ip = request.client.host if request.client else "unknown"
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    conn = get_db()
    try:
        user = conn.execute("SELECT id, password_hash FROM users WHERE username = ?", (username,)).fetchone()
        if not user:
            return UserAuthResponse(success=False, error="用户名不存在")

        if user["password_hash"] != hash_password(password):
            return UserAuthResponse(success=False, error="密码错误")

        # 更新 IP 和最后登录时间
        conn.execute(
            "UPDATE users SET ip_address = ?, last_login = ? WHERE id = ?",
            (client_ip, now, user["id"])
        )
        conn.commit()

        token = create_user_token(user["id"], username)
        return UserAuthResponse(success=True, token=token, username=username)
    finally:
        conn.close()


@app.get("/api/auth/verify")
async def user_verify(user: dict = Depends(user_required)):
    """验证用户 token 是否有效"""
    return {"success": True, "username": user["username"]}


# ============================================================
# 云端数据同步 API
# ============================================================

class SyncDataRequest(BaseModel):
    data: dict  # 前端全部 localStorage 数据（JSON）


@app.get("/api/sync/pull")
async def sync_pull(user: dict = Depends(user_required)):
    """从云端拉取用户数据"""
    user_id = user["user_id"]
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT data, updated_at FROM user_data WHERE user_id = ?", (user_id,)
        ).fetchone()
        if row:
            data = json.loads(row["data"])
            return {
                "success": True,
                "data": data,
                "updated_at": row["updated_at"],
                "empty": False
            }
        else:
            return {"success": True, "data": {}, "updated_at": None, "empty": True}
    finally:
        conn.close()


@app.post("/api/sync/push")
async def sync_push(req: SyncDataRequest, user: dict = Depends(user_required)):
    """将用户数据推送至云端"""
    user_id = user["user_id"]
    data_str = json.dumps(req.data, ensure_ascii=False)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    conn = get_db()
    try:
        conn.execute("""
            INSERT INTO user_data (user_id, data, updated_at)
            VALUES (?, ?, ?)
            ON CONFLICT(user_id) DO UPDATE SET
                data = excluded.data,
                updated_at = excluded.updated_at
        """, (user_id, data_str, now))
        conn.commit()
        return {"success": True, "updated_at": now}
    except Exception as e:
        raise HTTPException(500, f"同步失败: {str(e)}")
    finally:
        conn.close()


# ============================================================
# 用户端 API
# ============================================================

@app.post("/api/parse", response_model=ParseResponse)
async def parse_document(file: UploadFile = File(...), user: dict = Depends(user_required)):
    """
    用户上传文档，服务端解析题目（需登录）
    支持 .pdf / .docx / .doc / .txt
    """
    user_id = user["user_id"]
    username = user["username"]

    try:
        file_bytes = await file.read()
        filename = file.filename or "upload.doc"
        ext = os.path.splitext(filename)[1].lower()

        # 保存原始文件供管理员下载
        saved_file_path = ""
        uid_dir = os.path.join(UPLOAD_DIR, str(user_id))
        os.makedirs(uid_dir, exist_ok=True)
        safe_name = f"{int(time.time())}_{filename}"
        full_path = os.path.join(uid_dir, safe_name)
        with open(full_path, "wb") as f:
            f.write(file_bytes)
        saved_file_path = full_path

        # 1. 提取文本：优先使用 markitdown skill（如果启用）
        raw_text = None
        skill_used = None
        if is_skill_enabled("markitdown"):
            raw_text = convert_with_skill(file_bytes, filename)
            if raw_text:
                skill_used = "markitdown"

        if not raw_text:
            raw_text = extract_text(file_bytes, filename)

        if not raw_text or len(raw_text.strip()) < 10:
            add_log(filename, ext, False, 0, [], "未能提取有效文本", user_id, username, raw_text, saved_file_path)
            return ParseResponse(
                success=False,
                error="未能从文档中提取到有效文本，请检查文件格式"
            )

        # 2. 大模型解析
        questions = await parse_with_llm(raw_text)

        # 3. 过滤无效题目（只有符号没有实际内容的）
        questions = [q for q in questions if q.get("question", "").strip() and len(q.get("question", "").strip()) > 2 and q.get("question", "").strip() not in ("×", "√", "×", "√", "对", "错")]

        # 4. 标准化题型
        for q in questions:
            t = q.get("type", "s")
            if t == "j":
                q["options"] = ["A. 对", "B. 错"]
                # 标准化判断题答案：对/√/正确 → A, 错/×/错误 → B
                ans = str(q.get("answer", "")).strip()
                if ans in ("对", "√", "正确", "是", "A", "a"):
                    q["answer"] = "A"
                elif ans in ("错", "×", "错误", "否", "B", "b"):
                    q["answer"] = "B"
            elif t == "e":
                q["options"] = []
            q.setdefault("options", [])
            q.setdefault("answer", "")

        add_log(filename, ext, True, len(questions), questions, "", user_id, username, raw_text, saved_file_path)
        return ParseResponse(
            success=True,
            questions=questions,
            total=len(questions)
        )

    except HTTPException as e:
        add_log(filename, ext, False, 0, [], str(e.detail), user_id, username, raw_text if 'raw_text' in dir() else "", saved_file_path if 'saved_file_path' in dir() else "")
        raise
    except ValueError as e:
        add_log(file.filename or "unknown", "", False, 0, [], str(e), user_id, username)
        raise HTTPException(400, str(e))
    except Exception as e:
        add_log(file.filename or "unknown", "", False, 0, [], str(e), user_id, username)
        raise HTTPException(500, f"解析失败: {str(e)}")


class ParseTextRequest(BaseModel):
    text: str

@app.post("/api/parse-text", response_model=ParseResponse)
async def parse_text(req: ParseTextRequest, user: dict = Depends(user_required)):
    """
    用户粘贴文本，服务端解析题目（需登录）
    """
    user_id = user["user_id"]
    username = user["username"]

    try:
        text = req.text.strip()
        if not text or len(text) < 10:
            return ParseResponse(success=False, error="文本内容太少，无法解析")

        questions = await parse_with_llm(text)

        for q in questions:
            t = q.get("type", "s")
            if t == "j":
                q["options"] = ["A. 对", "B. 错"]
            elif t == "e":
                q["options"] = []
            q.setdefault("options", [])
            q.setdefault("answer", "")

        add_log("paste_text", "text", True, len(questions), questions, "", user_id, username, text)
        return ParseResponse(success=True, questions=questions, total=len(questions))

    except HTTPException as e:
        add_log("paste_text", "text", False, 0, [], str(e.detail), user_id, username, text)
        raise
    except Exception as e:
        add_log("paste_text", "text", False, 0, [], str(e), user_id, username)
        raise HTTPException(500, f"解析失败: {str(e)}")


@app.get("/api/health")
async def health():
    cfg = get_config()
    return {
        "status": "ok",
        "model": cfg["model"],
        "api_configured": bool(cfg["api_key"])
    }


# ============================================================
# 管理端 API
# ============================================================

@app.post("/api/admin/login", response_model=LoginResponse)
async def admin_login(req: LoginRequest):
    """管理员登录"""
    cfg = get_config()
    if req.password == cfg["admin_password"]:
        token = create_admin_token()
        return LoginResponse(success=True, token=token)
    return LoginResponse(success=False, error="密码错误")


@app.get("/api/admin/verify")
async def admin_verify(_: bool = Depends(admin_required)):
    """验证 token 是否有效"""
    return {"success": True}


@app.get("/api/admin/config")
async def admin_get_config(_: bool = Depends(admin_required)):
    """获取当前配置"""
    cfg = get_config()
    key = cfg["api_key"]
    masked = key[:4] + "****" + key[-4:] if len(key) > 8 else ("****" if key else "")
    return {
        "api_key_masked": masked,
        "api_base": cfg["api_base"],
        "model": cfg["model"],
        "api_configured": bool(key)
    }


@app.post("/api/admin/config")
async def admin_update_config(update: ConfigUpdate, _: bool = Depends(admin_required)):
    """更新配置"""
    cfg = get_config()
    changed = False

    if update.api_key and "*" not in update.api_key:
        cfg["api_key"] = update.api_key
        changed = True
    if update.api_base and update.api_base.strip():
        cfg["api_base"] = update.api_base.strip()
        changed = True
    if update.model and update.model.strip():
        model = update.model.strip()
        # 自动检测并切换 DeepSeek flash → pro
        if "deepseek" in model.lower() and "flash" in model.lower():
            model = "deepseek-chat"
        cfg["model"] = model
        changed = True

    if changed:
        save_config(cfg)
        # 更新全局配置
        global _config
        _config = cfg

    return {"success": True}


@app.post("/api/admin/change-password")
async def admin_change_password(req: ChangePasswordRequest, _: bool = Depends(admin_required)):
    """修改管理员密码（需验证旧密码 + 二次确认）"""
    cfg = get_config()

    if not req.old_password:
        raise HTTPException(400, "请输入旧密码")
    if req.old_password != cfg["admin_password"]:
        raise HTTPException(400, "旧密码错误")
    if not req.new_password or len(req.new_password.strip()) < 4:
        raise HTTPException(400, "新密码至少需要 4 位")
    if req.new_password != req.confirm_password:
        raise HTTPException(400, "两次输入的新密码不一致")
    if req.new_password == req.old_password:
        raise HTTPException(400, "新密码不能与旧密码相同")

    cfg["admin_password"] = req.new_password.strip()
    save_config(cfg)
    global _config
    _config = cfg
    return {"success": True, "message": "密码修改成功"}


@app.get("/api/admin/stats")
async def admin_stats(_: bool = Depends(admin_required)):
    """获取统计数据"""
    cfg = get_config()
    today = datetime.now().strftime("%Y-%m-%d")

    today_uploads = sum(1 for l in _logs if l["time"].startswith(today))
    today_questions = sum(l["question_count"] for l in _logs if l["time"].startswith(today) and l["success"])
    total_questions = sum(l["question_count"] for l in _logs if l["success"])

    recent = _logs[-10:][::-1]  # 最近10条，倒序

    return {
        "api_configured": bool(cfg["api_key"]),
        "model": cfg["model"],
        "today_uploads": today_uploads,
        "today_questions": today_questions,
        "total_questions": total_questions,
        "recent_logs": [
            {
                "time": l["time"],
                "filename": l["filename"],
                "question_count": l["question_count"],
                "success": l["success"],
                "username": l.get("username", "")
            }
            for l in recent
        ]
    }


@app.get("/api/admin/logs")
async def admin_logs(limit: int = 50, _: bool = Depends(admin_required)):
    """获取上传日志"""
    logs = _logs[-limit:][::-1]
    total = len(_logs)
    return {
        "logs": [
            {
                "time": l["time"],
                "filename": l["filename"],
                "format": l["format"],
                "success": l["success"],
                "question_count": l["question_count"],
                "questions": l.get("questions", []),
                "error": l.get("error", ""),
                "user_id": l.get("user_id", 0),
                "username": l.get("username", ""),
                "raw_text": l.get("raw_text", ""),
                "saved_file_path": l.get("saved_file_path", ""),
                "log_index": total - limit + i  # 全局索引，用于下载
            }
            for i, l in enumerate(logs)
        ]
    }


@app.get("/api/admin/skills")
async def admin_skills(_: bool = Depends(admin_required)):
    """获取所有注册的 Skills 及其状态"""
    import os as _os
    skills_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "skills")
    skills = []
    if os.path.exists(skills_dir):
        for name in os.listdir(skills_dir):
            skill_path = os.path.join(skills_dir, name)
            if os.path.isdir(skill_path):
                desc = ""
                skill_json = os.path.join(skill_path, "skill.json")
                if os.path.exists(skill_json):
                    try:
                        with open(skill_json, "r", encoding="utf-8") as f:
                            sj = json.load(f)
                        desc = sj.get("description", "")
                    except:
                        pass
                cfg = load_config()
                disabled = cfg.get("disabled_skills", [])
                skills.append({
                    "name": name,
                    "description": desc,
                    "enabled": name not in disabled
                })
    return {"success": True, "skills": skills}


@app.post("/api/admin/skills/toggle")
async def admin_skills_toggle(req: Request, _: bool = Depends(admin_required)):
    """启用/禁用某个 Skill"""
    import os as _os
    data = await req.json()
    name = data.get("name", "")
    enabled = data.get("enabled", True)
    if not name:
        raise HTTPException(400, "Skill 名称不能为空")

    cfg = load_config()
    disabled = cfg.get("disabled_skills", [])
    if enabled:
        if name in disabled:
            disabled.remove(name)
    else:
        if name not in disabled:
            disabled.append(name)
    cfg["disabled_skills"] = disabled

    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)
    return {"success": True, "enabled": enabled}


@app.get("/api/admin/users")
async def admin_users(_: bool = Depends(admin_required)):
    """获取所有注册用户列表"""
    conn = get_db()
    try:
        rows = conn.execute(
            "SELECT id, username, password_hash, ip_address, last_login, created_at FROM users ORDER BY id DESC"
        ).fetchall()
        users = []
        for r in rows:
            users.append({
                "id": r["id"],
                "username": r["username"],
                "password_hash": r["password_hash"],
                "ip_address": r["ip_address"],
                "last_login": r["last_login"],
                "created_at": r["created_at"]
            })
        return {"success": True, "users": users, "total": len(users)}
    finally:
        conn.close()


@app.delete("/api/admin/users/{user_id}")
async def admin_delete_user(user_id: int, _: bool = Depends(admin_required)):
    """删除用户及其同步数据"""
    conn = get_db()
    try:
        # 删除用户数据
        conn.execute("DELETE FROM user_data WHERE user_id = ?", (user_id,))
        # 删除用户
        cur = conn.execute("DELETE FROM users WHERE id = ?", (user_id,))
        conn.commit()
        if cur.rowcount == 0:
            raise HTTPException(404, "用户不存在")
        return {"success": True, "message": "用户已删除"}
    finally:
        conn.close()


class ResetPasswordRequest(BaseModel):
    new_password: str

@app.post("/api/admin/users/{user_id}/reset-password")
async def admin_reset_user_password(user_id: int, req: ResetPasswordRequest, _: bool = Depends(admin_required)):
    """管理员重置用户密码"""
    if not req.new_password or len(req.new_password.strip()) < 4:
        raise HTTPException(400, "新密码至少需要 4 位")
    conn = get_db()
    try:
        user = conn.execute("SELECT id FROM users WHERE id = ?", (user_id,)).fetchone()
        if not user:
            raise HTTPException(404, "用户不存在")
        new_hash = hash_password(req.new_password.strip())
        conn.execute("UPDATE users SET password_hash = ? WHERE id = ?", (new_hash, user_id))
        conn.commit()
        return {"success": True, "message": "密码已重置"}
    finally:
        conn.close()


# ============================================================
# 文件存储（用于管理员下载原始文件）
# ============================================================
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

@app.get("/api/admin/download/{log_index}")
async def admin_download_file(log_index: int, _: bool = Depends(admin_required)):
    """管理员下载用户上传的原始文件"""
    if log_index < 0 or log_index >= len(_logs):
        raise HTTPException(404, "记录不存在")
    log_entry = _logs[log_index]
    saved_path = log_entry.get("saved_file_path", "")
    if not saved_path or not os.path.exists(saved_path):
        raise HTTPException(404, "源文件不存在或已过期")
    from fastapi.responses import FileResponse
    return FileResponse(saved_path, filename=log_entry.get("filename", "download"))


@app.get("/api/admin/download-json/{log_index}")
async def admin_download_json(log_index: int, _: bool = Depends(admin_required)):
    """管理员下载解析后的 JSON 题目文件"""
    if log_index < 0 or log_index >= len(_logs):
        raise HTTPException(404, "记录不存在")
    log_entry = _logs[log_index]
    questions = log_entry.get("questions", [])
    if not questions:
        raise HTTPException(404, "该记录没有解析出题目")
    from fastapi.responses import Response
    json_str = json.dumps(questions, ensure_ascii=False, indent=2)
    filename = log_entry.get("filename", "questions")
    # 替换扩展名为 .json
    base = os.path.splitext(filename)[0]
    return Response(
        content=json_str,
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{base}.json"'}
    )


# ============================================================
# 页面路由
# ============================================================

@app.get("/admin")
async def admin_page():
    """管理后台页面"""
    admin_path = os.path.join(BASE_DIR, "admin.html")
    if os.path.exists(admin_path):
        with open(admin_path, encoding="utf-8") as f:
            return HTMLResponse(f.read())
    return HTMLResponse("<h1>管理后台文件不存在</h1>")


@app.get("/")
async def index():
    """用户端刷题页面"""
    demo_path = os.path.join(BASE_DIR, "刷题通demo.html")
    if os.path.exists(demo_path):
        with open(demo_path, encoding="utf-8") as f:
            return HTMLResponse(f.read())
    return HTMLResponse("<h1>刷题通 - 后端服务运行中</h1>")


if __name__ == "__main__":
    import uvicorn
    print(f"  管理后台: http://localhost:8800/admin")
    print(f"  用户端:   http://localhost:8800")
    print(f"  默认管理员密码: admin123 (可通过 ADMIN_PASSWORD 环境变量修改)")
    uvicorn.run(app, host="0.0.0.0", port=8800)