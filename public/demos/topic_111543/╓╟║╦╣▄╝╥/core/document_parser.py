"""文档解析器 - 支持多格式文档解析和目录扫描"""
import os
from typing import List, Dict

class DocumentParser:
    def __init__(self):
        self.supported_formats = ['.docx', '.pdf', '.xlsx', '.xls', '.txt', '.md']

    def parse_file(self, file_path: str) -> Dict:
        """根据文件扩展名调用相应的解析方法"""
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
        result = {'filename': filename, 'filepath': file_path, 'format': ext.lstrip('.'), 'content': ''}

        if ext in ['.txt', '.md']:
            return self._parse_text(file_path, result)
        elif ext in ['.xlsx', '.xls']:
            return self._parse_excel(file_path, result)
        elif ext == '.docx':
            return self._parse_word(file_path, result)
        elif ext == '.pdf':
            return self._parse_pdf(file_path, result)
        else:
            result['content'] = f'[不支持的文件格式: {ext}]'
            return result

    def _parse_text(self, file_path: str, result: Dict) -> Dict:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                result['content'] = f.read()
            result['info'] = f'{len(result["content"])}字符'
        except Exception as e:
            result['content'] = f'[解析错误: {str(e)}]'
        return result

    def _parse_excel(self, file_path: str, result: Dict) -> Dict:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            texts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                texts.append(f'=== Sheet: {sheet_name} ===')
                for row in ws.iter_rows(max_row=100, values_only=True):
                    vals = [str(c) for c in row if c is not None]
                    if vals:
                        texts.append(' | '.join(vals))
            result['content'] = '\n'.join(texts)
            result['info'] = f'{len(wb.sheetnames)} sheets'
            wb.close()
        except Exception as e:
            result['content'] = f'[Excel解析错误: {str(e)}]'
        return result

    def _parse_word(self, file_path: str, result: Dict) -> Dict:
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    tables.append(' | '.join(cells))
            result['content'] = '\n'.join(paragraphs + tables)
            result['info'] = f'{len(paragraphs)}段落'
        except Exception as e:
            result['content'] = f'[Word解析错误: {str(e)}]'
        return result

    def _parse_pdf(self, file_path: str, result: Dict) -> Dict:
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            texts = []
            for page_num in range(min(len(doc), 20)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    texts.append(f'[第{page_num+1}页]\n{text}')
            result['content'] = '\n'.join(texts)
            result['info'] = f'{len(doc)}页'
            doc.close()
        except Exception as e:
            result['content'] = f'[PDF解析错误: {str(e)}]'
        return result

    def scan_directory(self, directory_path: str) -> List[Dict]:
        """扫描目录中的文档文件，返回文件信息列表"""
        supported_extensions = ['.docx', '.doc', '.xlsx', '.xls', '.pdf', '.txt', '.md']
        documents = []
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in supported_extensions:
                    filepath = os.path.join(root, file)
                    size = os.path.getsize(filepath) if os.path.exists(filepath) else 0
                    documents.append({
                        'name': file,
                        'path': filepath,
                        'ext': ext,
                        'size': size,
                        'status': '已扫描'
                    })
        # 按文件大小排序
        documents.sort(key=lambda x: x['size'])
        return documents