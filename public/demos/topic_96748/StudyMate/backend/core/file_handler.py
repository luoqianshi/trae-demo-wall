"""
文件处理器 - 支持 Word 和图片文件的读取
"""

import base64
import io
import os
import re
from typing import Optional, Tuple, List

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


class FileHandler:

    IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}
    WORD_EXTS = {".docx", ".doc"}

    @classmethod
    def is_supported(cls, filepath: str) -> bool:
        ext = os.path.splitext(filepath)[1].lower()
        return ext in cls.IMAGE_EXTS or ext in cls.WORD_EXTS

    @classmethod
    def is_image(cls, filepath: str) -> bool:
        ext = os.path.splitext(filepath)[1].lower()
        return ext in cls.IMAGE_EXTS

    @classmethod
    def is_word(cls, filepath: str) -> bool:
        ext = os.path.splitext(filepath)[1].lower()
        return ext in cls.WORD_EXTS

    @classmethod
    def read_word(cls, filepath: str) -> str:
        if not HAS_DOCX:
            raise Exception("未安装 python-docx，无法读取 Word 文件。请运行: pip install python-docx")

        try:
            doc = DocxDocument(filepath)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        row_text = " | ".join(cells)
                        paragraphs.append(f"[表格行] {row_text}")

            content = "\n".join(paragraphs)

            if not content.strip():
                raise Exception("Word 文件内容为空或无法解析")

            return content

        except Exception as e:
            if "未安装" in str(e):
                raise
            raise Exception(f"读取 Word 文件失败: {e}")

    @classmethod
    def read_image_base64(cls, filepath: str) -> str:
        if not HAS_PIL:
            raise Exception("未安装 Pillow，无法处理图片。请运行: pip install Pillow")

        ext = os.path.splitext(filepath)[1].lower()

        mime_map = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
            ".webp": "image/webp"
        }
        mime_type = mime_map.get(ext, "image/png")

        try:
            with Image.open(filepath) as img:
                img = cls._resize_image(img, max_size=2048)

                buffer = io.BytesIO()
                if ext in {".jpg", ".jpeg"}:
                    img.save(buffer, format="JPEG", quality=85)
                else:
                    img.save(buffer, format="PNG")
                buffer.seek(0)

                b64_data = base64.b64encode(buffer.read()).decode("utf-8")
                return f"data:{mime_type};base64,{b64_data}"

        except Exception as e:
            raise Exception(f"读取图片失败: {e}")

    @classmethod
    def _resize_image(cls, img: Image.Image, max_size: int = 2048) -> Image.Image:
        width, height = img.size
        if width <= max_size and height <= max_size:
            return img

        if width > height:
            new_width = max_size
            new_height = int(height * (max_size / width))
        else:
            new_height = max_size
            new_width = int(width * (max_size / height))

        return img.resize((new_width, new_height), Image.LANCZOS)

    @classmethod
    def get_file_info(cls, filepath: str) -> dict:
        size = os.path.getsize(filepath)
        ext = os.path.splitext(filepath)[1].lower()
        name = os.path.basename(filepath)

        if size < 1024:
            size_str = f"{size} B"
        elif size < 1024 * 1024:
            size_str = f"{size / 1024:.1f} KB"
        else:
            size_str = f"{size / (1024 * 1024):.1f} MB"

        return {
            "name": name,
            "size": size,
            "size_str": size_str,
            "type": "图片" if ext in cls.IMAGE_EXTS else "Word文档",
            "ext": ext
        }


class MultimodalBuilder:

    @staticmethod
    def build_text_message(role: str, content: str) -> dict:
        return {"role": role, "content": content}

    @staticmethod
    def build_text_with_image_message(role: str, text: str, image_base64: str) -> dict:
        return {
            "role": role,
            "content": [
                {"type": "text", "text": text},
                {
                    "type": "image_url",
                    "image_url": {"url": image_base64}
                }
            ]
        }

    @staticmethod
    def build_multi_image_message(role: str, text: str, image_base64_list: List[str]) -> dict:
        content = [{"type": "text", "text": text}]
        for b64 in image_base64_list:
            content.append({
                "type": "image_url",
                "image_url": {"url": b64}
            })
        return {"role": role, "content": content}

    @staticmethod
    def is_multimodal_message(msg: dict) -> bool:
        content = msg.get("content", "")
        if isinstance(content, list):
            return True
        return False


if __name__ == "__main__":
    print("FileHandler 支持的文件类型:")
    print("  Word:", FileHandler.WORD_EXTS)
    print("  图片:", FileHandler.IMAGE_EXTS)